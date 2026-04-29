from __future__ import annotations

import base64
import json
import mimetypes
import os
import re
import shutil
import socket
import subprocess
import sys
import tempfile
import time
import webbrowser
from urllib import error as urlerror
from urllib import request as urlrequest
from urllib.parse import urlparse
from pathlib import Path
from threading import Thread
from typing import Any, Dict, List, Optional, Tuple

import gradio as gr
from PIL import Image

try:
    import fitz  # PyMuPDF
except Exception:
    fitz = None

try:
    from docx import Document
    from docx.shared import Inches
except Exception:
    Document = None
    Inches = None



# =========================================================
# Paths and config
# =========================================================

def app_base_dir() -> Path:
    """Return the folder containing this .py file or the packaged .exe.

    PyInstaller note:
    - In normal Python mode, use the script folder.
    - In packaged mode, use the executable folder, not sys._MEIPASS, because
      questions.json, images/, exports/, and models/ should remain editable
      external resources next to the exe.
    """
    if getattr(sys, "frozen", False):
        return Path(sys.executable).resolve().parent
    return Path(__file__).resolve().parent


APP_DIR = app_base_dir()
DATA_FILE = APP_DIR / "questions.json"
IMAGE_DIR = APP_DIR / "images"
EXPORT_DIR = APP_DIR / "exports"
IMAGE_DIR.mkdir(exist_ok=True)
EXPORT_DIR.mkdir(exist_ok=True)

# llama-server mode:
#   This Gradio app no longer loads Transformers/PyTorch model weights in-process.
#   Start llama-server separately, then point the UI to its OpenAI-compatible endpoint.
#
# Example for Qwen3-VL-2B-Instruct GGUF:
#   llama-server ^
#     -m D:/models/Qwen3-VL-2B-Instruct-GGUF/Qwen3VL-2B-Instruct-Q4_K_M.gguf ^
#     --mmproj D:/models/Qwen3-VL-2B-Instruct-GGUF/mmproj-Qwen3VL-2B-Instruct-F16.gguf ^
#     --host 127.0.0.1 --port 8080 -ngl 99 -c 8192 --jinja
#
# Example for Qwen3.5-0.8B GGUF:
#   llama-server ^
#     -m D:/models/Qwen3.5-0.8B-GGUF/Qwen3.5-0.8B-Q4_K_M.gguf ^
#     --mmproj D:/models/Qwen3.5-0.8B-GGUF/mmproj-F16.gguf ^
#     --host 127.0.0.1 --port 8080 -ngl 99 -c 8192 --jinja ^
#     --chat-template-kwargs "{\"enable_thinking\":false}"
DEFAULT_LLAMA_SERVER_URL = os.environ.get("LLAMA_SERVER_URL", "http://127.0.0.1:8080")
DEFAULT_LLAMA_MODEL_NAME = os.environ.get("LLAMA_SERVER_MODEL", "Qwen3.5-0.8B-UD-Q8_K_XL")

# Your current local CPU model. The relative APP_DIR path works when this script is placed in F:\\OneDrive\\AI_proj.
DEFAULT_LOCAL_GGUF_MODEL = os.environ.get(
    "LLAMA_GGUF_MODEL",
    str(APP_DIR / "models" / "Qwen3.5-0.8B-UD-Q8_K_XL.gguf"),
)
# Absolute fallback for the path you provided.
if not Path(DEFAULT_LOCAL_GGUF_MODEL).exists():
    DEFAULT_LOCAL_GGUF_MODEL = r"F:\OneDrive\AI_proj\models\Qwen3.5-0.8B-UD-Q8_K_XL.gguf"

# Multimodal Qwen3.5 requires a separate mmproj file. Put mmproj-F16.gguf in the same models folder.
DEFAULT_LOCAL_MMPROJ = os.environ.get(
    "LLAMA_MMPROJ",
    str(Path(DEFAULT_LOCAL_GGUF_MODEL).with_name("mmproj-F16.gguf")),
)
DEFAULT_LLAMA_SERVER_EXE = os.environ.get(
    "LLAMA_SERVER_EXE",
    str(APP_DIR /"llama" / "llama-server.exe") if (APP_DIR / "llama" / "llama-server.exe").exists() else "llama-server.exe",
)
DEFAULT_LLAMA_CONTEXT = int(os.environ.get("LLAMA_CONTEXT", "4096"))
DEFAULT_LLAMA_THREADS = max(1, min(os.cpu_count() or 8, 12))

# Backend/API cache and remote API defaults.
CONFIG_FILE = APP_DIR / "qsaver_backend_config.json"
BACKEND_LOCAL = "本地 llama-server"
BACKEND_API = "远程 API"
REGION_BASE_URLS = {
    "China / Beijing": "https://dashscope.aliyuncs.com/compatible-mode/v1",
    "Singapore / International": "https://dashscope-intl.aliyuncs.com/compatible-mode/v1",
    "US / Virginia": "https://dashscope-us.aliyuncs.com/compatible-mode/v1",
    "Hong Kong": "https://cn-hongkong.dashscope.aliyuncs.com/compatible-mode/v1",
}
DEFAULT_API_MODEL_NAME = os.environ.get("QSAVER_API_MODEL", "qwen3.6-max-preview")
DEFAULT_API_REGION = os.environ.get("QSAVER_API_REGION", "China / Beijing")
DEFAULT_API_BASE_URL = os.environ.get("QSAVER_API_BASE_URL", REGION_BASE_URLS[DEFAULT_API_REGION])

print(f"[App] APP_DIR = {APP_DIR}")
print(f"[App] DATA_FILE = {DATA_FILE}")
print(f"[App] IMAGE_DIR = {IMAGE_DIR}")
print(f"[App] EXPORT_DIR = {EXPORT_DIR}")
print(f"[App] DEFAULT_LLAMA_SERVER_URL = {DEFAULT_LLAMA_SERVER_URL}")
print(f"[App] DEFAULT_LOCAL_GGUF_MODEL = {DEFAULT_LOCAL_GGUF_MODEL}")
print(f"[App] DEFAULT_LOCAL_MMPROJ = {DEFAULT_LOCAL_MMPROJ}")

BIOLOGY_CATEGORIES = [
    "细胞生物学",
    "分子生物学",
    "遗传学",
    "生物化学",
    "动物生理学",
    "植物学",
    "微生物学",
    "免疫学",
    "发育生物学",
    "神经生物学",
    "生态学",
    "进化生物学",
    "脊椎动物学",
    "无脊椎动物学",
    "生物技术/实验技术",
    "文献题/图表题",
    "其他",
]

DEFAULT_SYSTEM_PROMPT = """你是一个严谨的生物学题库录入助手。
你的任务是从题目图片、题目文字、答案图片、解析文字中提取结构化题目信息。
原则：
1. 只基于输入中可见或明确提供的内容，不要凭空补全题干、选项、答案或解析。
2. OCR 看不清的字符用 <uncertain> 标记。
3. 数字、英文缩写、基因名、蛋白名、上下标符号尽量保持原样。
4. 如果无法确定答案，answer 留空字符串。
5. 分类可以根据题干与解析判断，但必须保守。
"""

CUSTOM_CSS = """
html, body { font-size: 18px !important; }
.gradio-container {
  max-width: 1880px !important;
  margin: 0 auto !important;
  padding: 12px 18px 18px 18px !important;
}
#topbar {
  position: relative;
  overflow: hidden;
  border-radius: 26px;
  padding: 22px 28px;
  margin-bottom: 16px;

  background:
    linear-gradient(
      135deg,
      rgba(255, 255, 255, 0.72),
      rgba(245, 247, 255, 0.46)
    );

  border: 1px solid rgba(255, 255, 255, 0.72);
  box-shadow:
    0 18px 45px rgba(31, 41, 55, 0.10),
    inset 0 1px 0 rgba(255, 255, 255, 0.85),
    inset 0 -1px 0 rgba(255, 255, 255, 0.35);

  backdrop-filter: blur(18px) saturate(160%);
  -webkit-backdrop-filter: blur(18px) saturate(160%);

  color: #1f2937;
}

#topbar::before {
  content: "";
  position: absolute;
  inset: 0;
  pointer-events: none;
  background:
    radial-gradient(
      circle at 12% 18%,
      rgba(124, 92, 255, 0.20),
      transparent 32%
    ),
    radial-gradient(
      circle at 88% 12%,
      rgba(56, 189, 248, 0.18),
      transparent 30%
    ),
    linear-gradient(
      120deg,
      rgba(255, 255, 255, 0.55),
      transparent 38%
    );
  opacity: 0.95;
}
#topbar::after {
  content: "";
  position: absolute;
  left: 20px;
  right: 20px;
  top: 10px;
  height: 1px;
  background: linear-gradient(
    90deg,
    transparent,
    rgba(255, 255, 255, 0.9),
    transparent
  );
}
#topbar h1 {
  position: relative;
  z-index: 1;
  margin: 0;
  font-size: 1.5rem;
  line-height: 1.15;
  letter-spacing: -0.035em;
  font-weight: 760;
  color: #111827;
}
#topbar p {
  position: relative;
  z-index: 1;
  margin: 9px 0 0 0;
  color: rgba(31, 41, 55, 0.68);
  font-size: 1rem;
  line-height: 1.55;
}
#left-panel, #right-panel, #center-panel {
  border: 1px solid #e6e8ef;
  border-radius: 22px;
  padding: 14px;
  min-height: 820px;
}
#left-panel { background: #f6f7fb; }
#center-panel { background: #ffffff; }
#right-panel { background: #f8f9fc; }
#q-input textarea, #a-input textarea, textarea, input { font-size: 1.02rem !important; line-height: 1.55 !important; }
button { font-size: 1rem !important; border-radius: 14px !important; }
label, .gradio-label, .block-label { font-size: 1rem !important; font-weight: 650 !important; }
.markdown, .prose, .gr-markdown { font-size: 1rem !important; line-height: 1.65 !important; }
#json-preview textarea { font-family: ui-monospace, SFMono-Regular, Menlo, Consolas, monospace !important; }
#db-table { max-height: 620px; overflow: auto; }
footer { display: none !important; }
"""


# =========================================================
# Persistent data
# =========================================================

def load_data() -> List[Dict[str, Any]]:
    if DATA_FILE.exists():
        with open(DATA_FILE, "r", encoding="utf-8") as f:
            data = json.load(f)
        if isinstance(data, list):
            return data
    return []


def save_data(data: List[Dict[str, Any]]) -> None:
    with open(DATA_FILE, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)


def append_question(q: Dict[str, Any]) -> int:
    data = load_data()
    data.append(q)
    save_data(data)
    return len(data)



# =========================================================
# Backend/API cache and model discovery helpers
# =========================================================

def load_backend_config() -> Dict[str, Any]:
    """Load user backend/API settings from a local JSON cache.

    The API key is saved in plain text because the user explicitly requested
    local caching. Keep the app folder private if you use this option.
    """
    if not CONFIG_FILE.exists():
        return {}
    try:
        data = json.loads(CONFIG_FILE.read_text(encoding="utf-8"))
        return data if isinstance(data, dict) else {}
    except Exception:
        return {}


def save_backend_config_dict(config: Dict[str, Any]) -> None:
    CONFIG_FILE.parent.mkdir(parents=True, exist_ok=True)
    CONFIG_FILE.write_text(json.dumps(config, ensure_ascii=False, indent=2), encoding="utf-8")


def config_get(config: Dict[str, Any], key: str, default: Any) -> Any:
    value = config.get(key, default)
    return default if value in (None, "") else value


def resolve_api_base_url(region: str, custom_base_url: str) -> str:
    custom = (custom_base_url or "").strip().rstrip("/")
    if custom:
        return custom
    return REGION_BASE_URLS.get(region, DEFAULT_API_BASE_URL).rstrip("/")


def is_mmproj_file(path: Path) -> bool:
    name = path.name.lower()
    return name.startswith("mmproj") or "mmproj" in name


def list_local_gguf_models() -> List[str]:
    """Return local .gguf main model files in ./models, excluding mmproj files."""
    models_dir = APP_DIR / "models"
    paths: List[Path] = []
    if models_dir.exists():
        paths.extend([p for p in models_dir.glob("*.gguf") if p.is_file() and not is_mmproj_file(p)])
    default_p = Path(DEFAULT_LOCAL_GGUF_MODEL)
    if default_p.exists() and default_p not in paths and not is_mmproj_file(default_p):
        paths.insert(0, default_p)
    unique: List[Path] = []
    seen = set()
    for p in paths:
        rp = str(p.resolve()) if p.exists() else str(p)
        if rp not in seen:
            seen.add(rp)
            unique.append(p)
    unique.sort(key=lambda p: (0 if str(p) == str(default_p) else 1, p.name.lower()))
    return [str(p) for p in unique]


def list_local_mmproj_files() -> List[str]:
    models_dir = APP_DIR / "models"
    paths: List[Path] = []
    if models_dir.exists():
        paths.extend([p for p in models_dir.glob("*.gguf") if p.is_file() and is_mmproj_file(p)])
    default_p = Path(DEFAULT_LOCAL_MMPROJ)
    if default_p.exists() and default_p not in paths:
        paths.insert(0, default_p)
    unique: List[Path] = []
    seen = set()
    for p in paths:
        rp = str(p.resolve()) if p.exists() else str(p)
        if rp not in seen:
            seen.add(rp)
            unique.append(p)
    unique.sort(key=lambda p: (0 if str(p) == str(default_p) else 1, p.name.lower()))
    return [str(p) for p in unique]


def local_model_name_from_path(path: str) -> str:
    p = Path((path or "").strip())
    return p.stem if p.name else DEFAULT_LLAMA_MODEL_NAME


def infer_mmproj_for_model(model_path: str, current_mmproj: str = "") -> str:
    """Prefer an existing current mmproj, then mmproj-F16.gguf next to the model, then first scanned mmproj."""
    if current_mmproj and Path(current_mmproj).exists():
        return current_mmproj
    p = Path((model_path or DEFAULT_LOCAL_GGUF_MODEL).strip())
    candidate = p.with_name("mmproj-F16.gguf")
    if candidate.exists():
        return str(candidate)
    mmprojs = list_local_mmproj_files()
    return mmprojs[0] if mmprojs else str(candidate)


def refresh_local_model_choices() -> Tuple[Any, Any, str]:
    models = list_local_gguf_models()
    mmprojs = list_local_mmproj_files()
    msg = f"已扫描 {APP_DIR / 'models'}：主模型 {len(models)} 个，mmproj {len(mmprojs)} 个。"
    return (
        gr.update(choices=models, value=models[0] if models else DEFAULT_LOCAL_GGUF_MODEL),
        gr.update(choices=mmprojs, value=mmprojs[0] if mmprojs else DEFAULT_LOCAL_MMPROJ),
        msg,
    )


def apply_local_model_choice(model_choice: str, current_mmproj: str) -> Tuple[str, str, str]:
    model_path = (model_choice or DEFAULT_LOCAL_GGUF_MODEL).strip()
    model_name = local_model_name_from_path(model_path)
    mmproj = infer_mmproj_for_model(model_path, current_mmproj)
    return model_path, model_name, mmproj


def update_backend_visibility(backend_mode: str):
    is_api = backend_mode == BACKEND_API
    return (
        gr.update(visible=not is_api),
        gr.update(visible=is_api),
        "当前后端：远程 API。不会自动启动本地 llama-server。" if is_api else "当前后端：本地 llama-server。Analyze 时会自动检测/启动本地模型服务。",
    )


def save_backend_settings(
    backend_mode: str,
    api_region: str,
    api_base_url: str,
    api_model_name: str,
    api_key: str,
    save_api_key: bool,
    server_url: str,
    model_name: str,
    local_model_path: str,
    local_mmproj_path: str,
) -> str:
    config = load_backend_config()
    config.update({
        "backend_mode": backend_mode or BACKEND_LOCAL,
        "api_region": api_region or DEFAULT_API_REGION,
        "api_base_url": api_base_url or "",
        "api_model_name": api_model_name or DEFAULT_API_MODEL_NAME,
        "server_url": server_url or DEFAULT_LLAMA_SERVER_URL,
        "model_name": model_name or DEFAULT_LLAMA_MODEL_NAME,
        "local_model_path": local_model_path or DEFAULT_LOCAL_GGUF_MODEL,
        "local_mmproj_path": local_mmproj_path or DEFAULT_LOCAL_MMPROJ,
        "save_api_key": bool(save_api_key),
    })
    if save_api_key:
        config["api_key"] = api_key or ""
    else:
        config.pop("api_key", None)
    save_backend_config_dict(config)
    return f"已保存到本地缓存：{CONFIG_FILE}"


# =========================================================
# API profile helpers
# =========================================================

def _mask_key(key: str) -> str:
    key = (key or "").strip()
    if not key:
        return "no-key"
    if len(key) <= 8:
        return "*" * len(key)
    return key[:4] + "..." + key[-4:]


def _api_profile_label(profile: Dict[str, Any]) -> str:
    name = str(profile.get("name") or "API profile").strip()
    base = str(profile.get("base_url") or profile.get("effective_base_url") or "region-default").strip()
    model = str(profile.get("model_name") or DEFAULT_API_MODEL_NAME).strip()
    return f"{name} | {model} | {base}"


def get_api_profiles(config: Optional[Dict[str, Any]] = None) -> List[Dict[str, Any]]:
    cfg = config if isinstance(config, dict) else load_backend_config()
    profiles = cfg.get("api_profiles", [])
    if not isinstance(profiles, list):
        return []
    return [p for p in profiles if isinstance(p, dict)]


def api_profile_choices() -> List[str]:
    return [_api_profile_label(p) for p in get_api_profiles()]


def _profile_name_from_values(api_region: str, base_url: str, model_name: str) -> str:
    effective = resolve_api_base_url(api_region, base_url)
    try:
        host = urlparse(effective).hostname or effective
    except Exception:
        host = effective
    model = (model_name or DEFAULT_API_MODEL_NAME).strip()
    return f"{host} / {model}"


def upsert_api_profile(
    api_region: str,
    api_base_url: str,
    api_model_name: str,
    api_key: str,
    save_api_key: bool = True,
    profile_name: str = "",
) -> Tuple[List[str], str]:
    """Save or update a URL-key pair in qsaver_backend_config.json.

    If api_base_url is blank, the profile records the selected region endpoint but
    keeps the editable custom base_url blank so the default Aliyun/DashScope path
    remains recoverable.
    """
    config = load_backend_config()
    profiles = get_api_profiles(config)
    api_region = api_region or DEFAULT_API_REGION
    api_base_url = (api_base_url or "").strip().rstrip("/")
    effective_base_url = resolve_api_base_url(api_region, api_base_url)
    api_model_name = (api_model_name or DEFAULT_API_MODEL_NAME).strip()
    api_key = (api_key or "").strip()
    name = (profile_name or _profile_name_from_values(api_region, api_base_url, api_model_name)).strip()

    found = False
    for p in profiles:
        if (
            str(p.get("effective_base_url") or "").rstrip("/") == effective_base_url.rstrip("/")
            and str(p.get("model_name") or "") == api_model_name
        ):
            p.update({
                "name": name,
                "region": api_region,
                "base_url": api_base_url,
                "effective_base_url": effective_base_url,
                "model_name": api_model_name,
                "updated_at": time.strftime("%Y-%m-%d %H:%M:%S"),
            })
            if save_api_key:
                p["api_key"] = api_key
            else:
                p.pop("api_key", None)
            found = True
            break
    if not found:
        profile = {
            "name": name,
            "region": api_region,
            "base_url": api_base_url,
            "effective_base_url": effective_base_url,
            "model_name": api_model_name,
            "created_at": time.strftime("%Y-%m-%d %H:%M:%S"),
            "updated_at": time.strftime("%Y-%m-%d %H:%M:%S"),
        }
        if save_api_key:
            profile["api_key"] = api_key
        profiles.append(profile)

    config["api_profiles"] = profiles
    config.update({
        "backend_mode": BACKEND_API,
        "api_region": api_region,
        "api_base_url": api_base_url,
        "api_model_name": api_model_name,
        "save_api_key": bool(save_api_key),
    })
    if save_api_key:
        config["api_key"] = api_key
    save_backend_config_dict(config)
    return api_profile_choices(), f"已保存 API 配置：{name}\nURL={effective_base_url}\nmodel={api_model_name}\nkey={_mask_key(api_key)}"


def load_api_profile(profile_label: str) -> Tuple[Any, str, str, str, bool, str]:
    profiles = get_api_profiles()
    if not profile_label:
        return gr.update(value=None), "", DEFAULT_API_MODEL_NAME, "", True, "未选择 API profile。"
    for p in profiles:
        if _api_profile_label(p) == profile_label:
            region = str(p.get("region") or DEFAULT_API_REGION)
            base_url = str(p.get("base_url") or "")
            model = str(p.get("model_name") or DEFAULT_API_MODEL_NAME)
            key = str(p.get("api_key") or "")
            return (
                gr.update(value=region),
                base_url,
                model,
                key,
                bool(key),
                f"已载入：{_api_profile_label(p)}\nkey={_mask_key(key)}",
            )
    return gr.update(value=DEFAULT_API_REGION), "", DEFAULT_API_MODEL_NAME, "", True, "没有找到该 API profile，可能缓存已变更。"


def delete_api_profile(profile_label: str) -> Tuple[Any, str]:
    config = load_backend_config()
    profiles = get_api_profiles(config)
    before = len(profiles)
    profiles = [p for p in profiles if _api_profile_label(p) != profile_label]
    config["api_profiles"] = profiles
    save_backend_config_dict(config)
    choices = api_profile_choices()
    return gr.update(choices=choices, value=None), f"已删除 {before - len(profiles)} 个 API profile。"


def test_api_and_save_profile(
    api_region: str,
    api_base_url: str,
    api_model_name: str,
    api_key: str,
    save_api_key: bool,
) -> Tuple[Any, str]:
    """Run a minimal OpenAI-compatible request; save the URL-key pair only if it succeeds."""
    effective_url = resolve_api_base_url(api_region, api_base_url)
    key = (api_key or os.getenv("DASHSCOPE_API_KEY", "")).strip()
    if not key:
        return gr.update(choices=api_profile_choices()), "测试失败：API key 为空。"
    try:
        text = run_vl_once(
            server_url=effective_url,
            model_name=(api_model_name or DEFAULT_API_MODEL_NAME).strip(),
            api_key=key,
            prompt="请只回复 OK。",
            images=[],
            system_prompt="你是一个连通性测试助手。",
            max_new_tokens=32,
            temperature=0.0,
            top_p=1.0,
        )
    except Exception as e:
        return gr.update(choices=api_profile_choices()), f"测试失败，不保存：{type(e).__name__}: {e}"
    choices, msg = upsert_api_profile(
        api_region=api_region,
        api_base_url=api_base_url,
        api_model_name=api_model_name,
        api_key=key,
        save_api_key=save_api_key,
    )
    return gr.update(choices=choices, value=choices[-1] if choices else None), f"测试通过，已保存。模型回复：{text[:200]}\n\n{msg}"


BACKEND_CONFIG = load_backend_config()
LOCAL_MODEL_CHOICES = list_local_gguf_models()
LOCAL_MMPROJ_CHOICES = list_local_mmproj_files()
API_PROFILE_CHOICES = api_profile_choices()

# =========================================================
# llama-server helpers
# =========================================================

def normalize_llama_endpoint(server_url: str) -> str:
    """Return the /v1/chat/completions endpoint for llama-server."""
    url = (server_url or DEFAULT_LLAMA_SERVER_URL).strip().rstrip("/")
    if not url:
        url = DEFAULT_LLAMA_SERVER_URL.rstrip("/")
    if url.endswith("/v1/chat/completions"):
        return url
    if url.endswith("/v1"):
        return url + "/chat/completions"
    return url + "/v1/chat/completions"


def image_to_data_url(image_path: str) -> str:
    """Encode a local image path as a data URL accepted by OpenAI-compatible multimodal APIs."""
    mime = mimetypes.guess_type(image_path)[0] or "image/png"
    with open(image_path, "rb") as f:
        b64 = base64.b64encode(f.read()).decode("ascii")
    return f"data:{mime};base64,{b64}"


def post_json(url: str, payload: Dict[str, Any], api_key: str = "", timeout: int = 600) -> Dict[str, Any]:
    """POST JSON with user-facing diagnostics for llama-server connection failures."""
    data = json.dumps(payload, ensure_ascii=False).encode("utf-8")
    headers = {"Content-Type": "application/json"}
    if api_key.strip():
        headers["Authorization"] = f"Bearer {api_key.strip()}"
    req = urlrequest.Request(url, data=data, headers=headers, method="POST")
    try:
        with urlrequest.urlopen(req, timeout=timeout) as resp:
            body = resp.read().decode("utf-8", errors="replace")
    except urlerror.HTTPError as e:
        err = e.read().decode("utf-8", errors="replace")
        raise RuntimeError(
            f"OpenAI-compatible 后端已响应，但请求失败：HTTP {e.code}\n"
            f"endpoint={url}\n"
            f"server_message={err[:2000]}"
        ) from e
    except urlerror.URLError as e:
        endpoint_base = url.rsplit("/v1/", 1)[0] if "/v1/" in url else url
        host, port = server_url_to_host_port(endpoint_base)
        port_msg = "端口已打开" if is_tcp_port_open(host, port, timeout=0.5) else "端口未监听"
        raise RuntimeError(
            "无法连接后端服务。\n"
            f"endpoint={url}\n"
            f"host={host}, port={port}, check={port_msg}\n"
            f"原始错误：{e}\n"
            "处理方式：本地模式请点击 Start local llama-server 或运行 Run diagnostics；API 模式请检查 base_url/key/网络。"
        ) from e
    try:
        return json.loads(body)
    except Exception as e:
        raise RuntimeError(
            "后端返回的不是合法 JSON。\n"
            f"endpoint={url}\n"
            f"response_head={body[:1000]}"
        ) from e

def extract_openai_message_content(resp: Dict[str, Any]) -> str:
    try:
        content = resp["choices"][0]["message"]["content"]
    except Exception as e:
        raise RuntimeError(f"后端返回格式异常：{json.dumps(resp, ensure_ascii=False)[:1200]}") from e

    if isinstance(content, str):
        return content
    if isinstance(content, list):
        parts = []
        for item in content:
            if isinstance(item, dict):
                parts.append(str(item.get("text") or item.get("content") or ""))
            else:
                parts.append(str(item))
        return "".join(parts)
    return str(content or "")


# Keep the subprocess alive while the Gradio app is running.
_LLAMA_SERVER_PROC: Optional[subprocess.Popen] = None


def server_url_to_host_port(server_url: str) -> Tuple[str, int]:
    """Parse host and port from a llama-server base URL."""
    parsed = urlparse((server_url or DEFAULT_LLAMA_SERVER_URL).strip())
    host = parsed.hostname or "127.0.0.1"
    port = parsed.port or (443 if parsed.scheme == "https" else 80)
    return host, int(port)


def is_tcp_port_open(host: str, port: int, timeout: float = 0.6) -> bool:
    try:
        with socket.create_connection((host, port), timeout=timeout):
            return True
    except OSError:
        return False


def http_get_json(url: str, api_key: str = "", timeout: float = 3.0) -> Tuple[bool, str, Optional[Dict[str, Any]]]:
    """GET JSON and return (ok, message, object). Does not raise."""
    headers: Dict[str, str] = {}
    if api_key.strip():
        headers["Authorization"] = f"Bearer {api_key.strip()}"
    req = urlrequest.Request(url, headers=headers, method="GET")
    try:
        with urlrequest.urlopen(req, timeout=timeout) as resp:
            body = resp.read().decode("utf-8", errors="replace")
        try:
            obj = json.loads(body)
        except Exception:
            return False, f"GET {url} 有响应，但不是 JSON：{body[:500]}", None
        return True, f"GET {url} 成功", obj
    except urlerror.HTTPError as e:
        body = e.read().decode("utf-8", errors="replace")
        return False, f"GET {url} HTTP {e.code}: {body[:500]}", None
    except Exception as e:
        return False, f"GET {url} 失败：{e}", None


def llama_base_url(server_url: str) -> str:
    url = (server_url or DEFAULT_LLAMA_SERVER_URL).strip().rstrip("/")
    if url.endswith("/v1/chat/completions"):
        return url[:-len("/v1/chat/completions")]
    if url.endswith("/v1"):
        return url[:-len("/v1")]
    return url


def llama_models_url(server_url: str) -> str:
    return llama_base_url(server_url) + "/v1/models"


def check_llama_server_ready(server_url: str, api_key: str = "", timeout: float = 2.0) -> Tuple[bool, str]:
    """Check whether an OpenAI-compatible llama-server endpoint is usable."""
    base = llama_base_url(server_url)
    host, port = server_url_to_host_port(base)
    if not is_tcp_port_open(host, port, timeout=0.5):
        return False, f"{host}:{port} 未监听。llama-server 尚未启动，或端口/URL 不一致。"
    ok, msg, obj = http_get_json(llama_models_url(base), api_key=api_key, timeout=timeout)
    if ok:
        model_info = ""
        try:
            data = obj.get("data", []) if isinstance(obj, dict) else []
            if data:
                names = [str(x.get("id", "")) for x in data[:3] if isinstance(x, dict)]
                model_info = "；models=" + ", ".join([x for x in names if x])
        except Exception:
            pass
        return True, f"llama-server 已就绪：{base}{model_info}"
    ok_root, msg_root, _ = http_get_json(base + "/health", api_key=api_key, timeout=timeout)
    if ok_root:
        return True, f"llama-server health endpoint 已响应：{base}"
    return False, msg + "\n" + msg_root


def file_size_text(path: str) -> str:
    try:
        pp = Path(path)
        if not pp.exists():
            return "missing"
        size = pp.stat().st_size
        if size >= 1024 ** 3:
            return f"{size / (1024 ** 3):.2f} GB"
        return f"{size / (1024 ** 2):.1f} MB"
    except Exception as e:
        return f"无法读取大小：{e}"


def tail_text(path: Path, max_lines: int = 80) -> str:
    if not path.exists():
        return ""
    try:
        lines = path.read_text(encoding="utf-8", errors="replace").splitlines()
        return "\n".join(lines[-max_lines:])
    except Exception as e:
        return f"无法读取日志：{e}"


def diagnose_llama_setup(
    llama_server_exe: str,
    model_path: str,
    mmproj_path: str,
    server_url: str,
    ctx_size: int,
    threads: int,
    api_key: str = "",
) -> str:
    """Return a detailed diagnostic report for paths, port, server readiness, and launch command."""
    llama_server_exe = (llama_server_exe or DEFAULT_LLAMA_SERVER_EXE).strip()
    model_path = (model_path or DEFAULT_LOCAL_GGUF_MODEL).strip()
    mmproj_path = (mmproj_path or DEFAULT_LOCAL_MMPROJ).strip()
    server_url = (server_url or DEFAULT_LLAMA_SERVER_URL).strip()
    ctx_size = int(ctx_size or DEFAULT_LLAMA_CONTEXT)
    threads = int(threads or DEFAULT_LLAMA_THREADS)
    lines: List[str] = []
    lines.append("[Runtime diagnostics]")
    lines.append(f"APP_DIR = {APP_DIR}")
    lines.append(f"Python = {sys.executable}")
    lines.append(f"CPU threads requested = {threads}; os.cpu_count = {os.cpu_count()}")
    lines.append("")
    exe_resolved = resolve_executable(llama_server_exe)
    if exe_resolved:
        lines.append(f"[OK] llama-server.exe = {exe_resolved}")
    else:
        lines.append(f"[ERROR] 找不到 llama-server.exe：{llama_server_exe}")
        lines.append("        请把 llama-server.exe 放到此脚本同目录，或在界面里填写完整路径。")
    model_p = Path(model_path)
    if model_p.exists():
        lines.append(f"[OK] 主模型 = {model_path} ({file_size_text(model_path)})")
        if model_p.suffix.lower() != ".gguf":
            lines.append("[WARN] 主模型后缀不是 .gguf，请确认没有选错文件。")
    else:
        lines.append(f"[ERROR] 找不到主模型文件：{model_path}")
    mmproj_p = Path(mmproj_path)
    if mmproj_path:
        if mmproj_p.exists():
            lines.append(f"[OK] mmproj = {mmproj_path} ({file_size_text(mmproj_path)})")
            if mmproj_p.suffix.lower() != ".gguf":
                lines.append("[WARN] mmproj 后缀不是 .gguf，请确认没有选错文件。")
        else:
            lines.append(f"[ERROR] 找不到 mmproj 文件：{mmproj_path}")
            lines.append("        图像输入必需 mmproj-F16.gguf；纯文本才可以不填。")
    else:
        lines.append("[WARN] 未填写 mmproj；有图片输入时会失败或无法识图。")
    host, port = server_url_to_host_port(server_url)
    tcp_open = is_tcp_port_open(host, port, timeout=0.5)
    lines.append(f"Server URL = {server_url}")
    lines.append(f"TCP {host}:{port} = {'OPEN' if tcp_open else 'CLOSED'}")
    ready, ready_msg = check_llama_server_ready(server_url, api_key=api_key, timeout=2.0)
    lines.append(f"Readiness = {'READY' if ready else 'NOT READY'}")
    lines.append(ready_msg)
    lines.append("")
    lines.append("[Launch command]")
    lines.append(format_llama_server_command(
        llama_server_exe=exe_resolved or llama_server_exe,
        model_path=model_path,
        mmproj_path=mmproj_path,
        server_url=server_url,
        ctx_size=ctx_size,
        threads=threads,
    ))
    log_path = APP_DIR / "logs" / "llama_server.log"
    if log_path.exists():
        lines.append("")
        lines.append(f"[Recent log tail] {log_path}")
        lines.append(tail_text(log_path, max_lines=50))
    return "\n".join(lines)



def resolve_executable(exe: str) -> Optional[str]:
    exe = (exe or "").strip()
    if not exe:
        return None
    p = Path(exe)
    if p.exists():
        return str(p)
    found = shutil.which(exe)
    return found


def quote_cmd_arg(arg: str) -> str:
    arg = str(arg)
    if not arg:
        return '""'
    if re.search(r'\s|["^&()<>|]', arg):
        return '"' + arg.replace('"', r'\"') + '"'
    return arg


def build_llama_server_cmd(
    llama_server_exe: str,
    model_path: str,
    mmproj_path: str,
    server_url: str,
    ctx_size: int,
    threads: int,
) -> List[str]:
    host, port = server_url_to_host_port(server_url)
    cmd = [
        llama_server_exe,
        "-m", model_path,
    ]
    if (mmproj_path or "").strip():
        cmd.extend(["--mmproj", mmproj_path])
    cmd.extend([
        "--host", host,
        "--port", str(port),
        "-c", str(int(ctx_size or DEFAULT_LLAMA_CONTEXT)),
        "-t", str(int(threads or DEFAULT_LLAMA_THREADS)),
        "--jinja",
        "--chat-template-kwargs", '{"enable_thinking":false}',
    ])
    return cmd


def format_llama_server_command(
    llama_server_exe: str,
    model_path: str,
    mmproj_path: str,
    server_url: str,
    ctx_size: int,
    threads: int,
) -> str:
    cmd = build_llama_server_cmd(
        llama_server_exe=llama_server_exe,
        model_path=model_path,
        mmproj_path=mmproj_path,
        server_url=server_url,
        ctx_size=ctx_size,
        threads=threads,
    )
    return " ".join(quote_cmd_arg(x) for x in cmd)


def start_local_llama_server(
    llama_server_exe: str,
    model_path: str,
    mmproj_path: str,
    server_url: str,
    ctx_size: int,
    threads: int,
) -> str:
    """Start llama-server as a subprocess with preflight checks and log diagnostics."""
    global _LLAMA_SERVER_PROC
    model_path = (model_path or DEFAULT_LOCAL_GGUF_MODEL).strip()
    mmproj_path = (mmproj_path or DEFAULT_LOCAL_MMPROJ).strip()
    llama_server_exe = (llama_server_exe or DEFAULT_LLAMA_SERVER_EXE).strip()
    server_url = (server_url or DEFAULT_LLAMA_SERVER_URL).strip()
    ctx_size = int(ctx_size or DEFAULT_LLAMA_CONTEXT)
    threads = int(threads or DEFAULT_LLAMA_THREADS)
    host, port = server_url_to_host_port(server_url)
    ready, ready_msg = check_llama_server_ready(server_url, timeout=2.0)
    if ready:
        return "已检测到可用的 llama-server。\n" + ready_msg
    if is_tcp_port_open(host, port, timeout=0.5):
        return (
            f"启动中止：{host}:{port} 已被占用，但不像可用的 llama-server。\n"
            f"检测结果：{ready_msg}\n\n"
            "处理方式：关闭占用 8080 的程序，或把 Server URL 改为 http://127.0.0.1:8081 后再启动。"
        )
    exe_resolved = resolve_executable(llama_server_exe)
    errors: List[str] = []
    if not exe_resolved:
        errors.append("找不到 llama-server.exe。请把 llama-server.exe 放到代码同目录，或把它加入 PATH，或在界面里填写完整路径。")
    if not Path(model_path).exists():
        errors.append(f"找不到主模型文件：{model_path}")
    if mmproj_path and not Path(mmproj_path).exists():
        errors.append(f"找不到 mmproj 文件：{mmproj_path}\nQwen3.5 图像输入需要单独的 mmproj-F16.gguf。")
    if errors:
        return "启动失败：\n- " + "\n- ".join(errors) + "\n\n" + diagnose_llama_setup(
            llama_server_exe, model_path, mmproj_path, server_url, ctx_size, threads
        )
    log_dir = APP_DIR / "logs"
    log_dir.mkdir(exist_ok=True)
    log_path = log_dir / "llama_server.log"
    cmd = build_llama_server_cmd(
        llama_server_exe=exe_resolved,
        model_path=model_path,
        mmproj_path=mmproj_path,
        server_url=server_url,
        ctx_size=ctx_size,
        threads=threads,
    )
    log_f = open(log_path, "a", encoding="utf-8", errors="replace")
    log_f.write("\n\n===== Starting llama-server at " + time.strftime("%Y-%m-%d %H:%M:%S") + " =====\n")
    log_f.write(" ".join(quote_cmd_arg(x) for x in cmd) + "\n")
    log_f.flush()
    try:
        _LLAMA_SERVER_PROC = subprocess.Popen(
            cmd,
            cwd=str(APP_DIR),
            stdout=log_f,
            stderr=subprocess.STDOUT,
            creationflags=getattr(subprocess, "CREATE_NEW_PROCESS_GROUP", 0),
        )
    except Exception as e:
        return f"启动失败：{e}\n\n" + diagnose_llama_setup(
            llama_server_exe, model_path, mmproj_path, server_url, ctx_size, threads
        )
    wait_s = int(os.environ.get("LLAMA_STARTUP_TIMEOUT", "90"))
    last_msg = ""
    for _ in range(wait_s):
        time.sleep(1.0)
        if _LLAMA_SERVER_PROC.poll() is not None:
            return (
                f"llama-server 启动后退出，退出码 {_LLAMA_SERVER_PROC.returncode}。\n"
                f"日志：{log_path}\n\n"
                "[log tail]\n" + tail_text(log_path, max_lines=100)
            )
        ready, last_msg = check_llama_server_ready(server_url, timeout=1.5)
        if ready:
            return (
                f"已启动并检测到 llama-server 可用。\n"
                f"地址：{server_url}\n"
                f"主模型：{model_path} ({file_size_text(model_path)})\n"
                f"mmproj：{mmproj_path} ({file_size_text(mmproj_path)})\n"
                f"线程数：{threads}，上下文：{ctx_size}\n"
                f"日志：{log_path}\n"
                f"检测：{last_msg}\n"
                "现在可以点击 Analyze。"
            )
    return (
        f"llama-server 进程仍在运行，但 {wait_s}s 内还没有通过 /v1/models 检测。\n"
        "可能仍在加载模型，也可能参数不兼容。请稍等后点击 Run diagnostics。\n"
        f"最后检测结果：{last_msg}\n"
        f"日志：{log_path}\n\n"
        "[log tail]\n" + tail_text(log_path, max_lines=100)
    )



def preview_llama_server_command(
    llama_server_exe: str,
    model_path: str,
    mmproj_path: str,
    server_url: str,
    ctx_size: int,
    threads: int,
) -> str:
    return format_llama_server_command(
        llama_server_exe=(llama_server_exe or DEFAULT_LLAMA_SERVER_EXE),
        model_path=(model_path or DEFAULT_LOCAL_GGUF_MODEL),
        mmproj_path=(mmproj_path or DEFAULT_LOCAL_MMPROJ),
        server_url=(server_url or DEFAULT_LLAMA_SERVER_URL),
        ctx_size=int(ctx_size or DEFAULT_LLAMA_CONTEXT),
        threads=int(threads or DEFAULT_LLAMA_THREADS),
    )


def run_llama_diagnostics(
    llama_server_exe: str,
    model_path: str,
    mmproj_path: str,
    server_url: str,
    ctx_size: int,
    threads: int,
    api_key: str = "",
) -> str:
    return diagnose_llama_setup(
        llama_server_exe=llama_server_exe,
        model_path=model_path,
        mmproj_path=mmproj_path,
        server_url=server_url,
        ctx_size=ctx_size,
        threads=threads,
        api_key=api_key,
    )


def ensure_llama_server_running(
    llama_server_exe: str,
    model_path: str,
    mmproj_path: str,
    server_url: str,
    ctx_size: int,
    threads: int,
    api_key: str = "",
) -> Tuple[bool, str]:
    """Ensure llama-server is ready before a model call.

    Clicking Analyze now automatically starts llama-server if the configured
    port is not already serving /v1/models.
    """
    server_url = (server_url or DEFAULT_LLAMA_SERVER_URL).strip()
    ready, msg = check_llama_server_ready(server_url, api_key=api_key, timeout=2.0)
    if ready:
        return True, "llama-server 已经可用。\n" + msg

    start_msg = start_local_llama_server(
        llama_server_exe=llama_server_exe,
        model_path=model_path,
        mmproj_path=mmproj_path,
        server_url=server_url,
        ctx_size=ctx_size,
        threads=threads,
    )
    ready2, msg2 = check_llama_server_ready(server_url, api_key=api_key, timeout=3.0)
    if ready2:
        return True, "自动启动 llama-server 成功。\n" + start_msg + "\n\n最终检测：" + msg2

    diag = diagnose_llama_setup(
        llama_server_exe=llama_server_exe,
        model_path=model_path,
        mmproj_path=mmproj_path,
        server_url=server_url,
        ctx_size=ctx_size,
        threads=threads,
        api_key=api_key,
    )
    return False, (
        "自动启动 llama-server 失败，或启动后未能通过 /v1/models 检测。\n\n"
        "[启动结果]\n" + start_msg + "\n\n[诊断]\n" + diag
    )


# =========================================================
# File and image helpers
# =========================================================

def now_stamp() -> str:
    return time.strftime("%Y%m%d_%H%M%S") + f"_{int(time.time() * 1000) % 1000:03d}"


def get_uploaded_file_path(file_obj: Any) -> Optional[str]:
    if file_obj is None:
        return None
    if isinstance(file_obj, str):
        return file_obj
    if isinstance(file_obj, dict):
        for key in ("path", "name", "orig_name"):
            value = file_obj.get(key)
            if value and os.path.exists(str(value)):
                return str(value)
    for attr in ("path", "name"):
        value = getattr(file_obj, attr, None)
        if value and os.path.exists(str(value)):
            return str(value)
    return None


def normalize_image(image_path: str, max_side: int = 1500) -> str:
    img = Image.open(image_path).convert("RGB")
    if max(img.size) > max_side:
        img.thumbnail((max_side, max_side))
    out = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
    img.save(out.name, format="PNG")
    return out.name


def pdf_to_images(pdf_path: str, max_pages: int = 2, zoom: float = 2.0, max_side: int = 1500) -> List[str]:
    if fitz is None:
        raise RuntimeError("PyMuPDF 未安装。请执行: pip install pymupdf")
    doc = fitz.open(pdf_path)
    paths: List[str] = []
    for i in range(min(len(doc), int(max_pages))):
        page = doc.load_page(i)
        pix = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom), alpha=False)
        raw = tempfile.NamedTemporaryFile(delete=False, suffix=f"_page_{i + 1}.png")
        pix.save(raw.name)
        paths.append(normalize_image(raw.name, max_side=max_side))
    doc.close()
    return paths


def files_to_images(uploaded_files: List[Any], max_pdf_pages: int, max_image_side: int) -> Tuple[List[str], List[str]]:
    image_paths: List[str] = []
    warnings: List[str] = []
    for file_obj in uploaded_files or []:
        p = get_uploaded_file_path(file_obj)
        if not p:
            warnings.append("有一个上传文件无法读取，已跳过。")
            continue
        suffix = Path(p).suffix.lower()
        try:
            if suffix == ".pdf":
                image_paths.extend(pdf_to_images(p, max_pages=max_pdf_pages, max_side=max_image_side))
            elif suffix in {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tiff", ".tif"}:
                image_paths.append(normalize_image(p, max_side=max_image_side))
            else:
                warnings.append(f"暂不支持文件类型：{Path(p).name}")
        except Exception as e:
            warnings.append(f"处理文件失败：{Path(p).name}: {e}")
    return image_paths, warnings


def save_question_image_for_json(question_images: List[str]) -> str:
    """The old schema has a single image field. Multiple images are stacked into one PNG."""
    if not question_images:
        return ""

    out_path = IMAGE_DIR / f"quiz_{now_stamp()}.png"

    if len(question_images) == 1:
        shutil.copy(question_images[0], out_path)
        return str(out_path.relative_to(APP_DIR))

    imgs = [Image.open(p).convert("RGB") for p in question_images]
    widths = [im.width for im in imgs]
    heights = [im.height for im in imgs]
    pad = 24
    total_w = max(widths)
    total_h = sum(heights) + pad * (len(imgs) - 1)
    canvas = Image.new("RGB", (total_w, total_h), "white")
    y = 0
    for im in imgs:
        canvas.paste(im, (0, y))
        y += im.height + pad
    canvas.save(out_path, format="PNG")
    return str(out_path.relative_to(APP_DIR))


# =========================================================
# JSON parsing and cleaning
# =========================================================

def extract_json_object(text: str) -> Dict[str, Any]:
    if not text:
        return {}
    cleaned = text.strip()
    cleaned = re.sub(r"^```(?:json)?", "", cleaned, flags=re.I).strip()
    cleaned = re.sub(r"```$", "", cleaned).strip()

    try:
        obj = json.loads(cleaned)
        return obj if isinstance(obj, dict) else {}
    except Exception:
        pass

    start = cleaned.find("{")
    end = cleaned.rfind("}")
    if start >= 0 and end > start:
        candidate = cleaned[start:end + 1]
        try:
            obj = json.loads(candidate)
            return obj if isinstance(obj, dict) else {}
        except Exception:
            return {}
    return {}


def clean_option_text(opt: str) -> str:
    opt = str(opt or "").strip()
    # remove A. / A、 / （A） / A: prefixes if the model included them
    opt = re.sub(r"^\s*[A-Ha-h][\.、:：\)]\s*", "", opt).strip()
    opt = re.sub(r"^\s*[（(][A-Ha-h][）)]\s*", "", opt).strip()
    return opt


def normalize_options(options: Any) -> List[str]:
    if isinstance(options, list):
        return [clean_option_text(x) for x in options if str(x).strip()]
    if isinstance(options, str):
        lines = [x.strip() for x in options.splitlines() if x.strip()]
        return [clean_option_text(x) for x in lines]
    return []


def normalize_tags(tags: Any, category: str = "") -> List[str]:
    out: List[str] = []
    if category:
        out.append(category.strip().lstrip("#"))
    if isinstance(tags, str):
        parts = re.split(r"[，,\s]+", tags)
    elif isinstance(tags, list):
        parts = [str(x) for x in tags]
    else:
        parts = []
    for p in parts:
        p = p.strip().lstrip("#")
        if p and p not in out:
            out.append(p)
    return out


def format_remark(tags: List[str], explanation: str,  extra_note: str = "") -> str:
    tag_line = " ".join(f"#{t}" for t in tags if t)
    blocks = []
    if tag_line:
        blocks.append(tag_line)
    if explanation.strip():
        blocks.append(explanation.strip())
    if extra_note.strip():
        blocks.append("[人工补充]\n" + extra_note.strip())
    # Do not put question OCR in remark by default; it can be very long and duplicates stem/options.
    return "\n".join(blocks).strip()


# =========================================================
# llama-server model calls
# =========================================================

def run_vl_once(
    server_url: str,
    model_name: str,
    api_key: str,
    prompt: str,
    images: List[str],
    system_prompt: str,
    max_new_tokens: int,
    temperature: float,
    top_p: float,
) -> str:
    """Call a running llama-server through its OpenAI-compatible chat endpoint."""
    endpoint = normalize_llama_endpoint(server_url)

    content: List[Dict[str, Any]] = []
    for img in images:
        content.append({
            "type": "image_url",
            "image_url": {"url": image_to_data_url(img)},
        })
    content.append({"type": "text", "text": prompt})

    payload: Dict[str, Any] = {
        "model": (model_name or DEFAULT_LLAMA_MODEL_NAME).strip() or DEFAULT_LLAMA_MODEL_NAME,
        "messages": [
            {"role": "system", "content": system_prompt.strip() or DEFAULT_SYSTEM_PROMPT},
            {"role": "user", "content": content},
        ],
        "max_tokens": int(max_new_tokens),
        "temperature": float(temperature),
        "top_p": float(top_p),
        "stream": False,
    }

    resp = post_json(endpoint, payload, api_key=api_key)
    return extract_openai_message_content(resp)


def build_question_prompt(question_text: str) -> str:
    cats = "、".join(BIOLOGY_CATEGORIES)
    return f"""
你将看到一道生物学选择/判断题的题干与选项，可能来自图片 OCR，也可能包含用户输入文字。
请完成：
1. 忠实提取题干 stem，去掉数字题号。
2. 忠实提取选项 options，去掉 A/B/C/D 标号，只保留选项文本。
3. 判断题目属于哪些生物学大类 category_tags。可从以下类别中选 1-3 个：{cats}。
4. 如果图片中包含题目图表，请题干中保留“如图/下图”等语义，不要虚构图中不存在的结论。
5. 输出严格 JSON，不要输出解释、Markdown 或代码块。

用户补充文字：
{question_text.strip() if question_text.strip() else "无"}

JSON schema:
{{
  "stem": "题干文本；不要包含选项",
  "options": ["选项A文本", "选项B文本", "选项C文本", "选项D文本"],
  "category_tags": ["细胞生物学"],
}}
""".strip()


def build_solution_prompt(solution_text: str, stem: str, options: List[str]) -> str:
    cats = "、".join(BIOLOGY_CATEGORIES)
    options_text = "\n".join(f"{chr(65+i)}. {opt}" for i, opt in enumerate(options))
    return f"""
你将看到答案/解析材料图片，也可能有用户输入文字。请继续忠实提取答案与解析。

用户补充的答案/解析文字：
{solution_text.strip() if solution_text.strip() else "无"}

要求：
1. answer 字段：忠实提取含T/F的字符串作为答案。如果用户输入了类似字段请使用用户输入。
2. explanation 字段：忠实提取图片中所有的文字。
3. category_tags 字段：从以下类别中选 1-3 个：{cats}。
4. 输出严格 JSON，不要输出解释、Markdown 或代码块。

JSON schema:
{{
  "answer": "TFTT",
  "explanation": "",
  "category_tags": ["细胞生物学"],
}}
""".strip()


# =========================================================
# Gradio callbacks
# =========================================================

def parse_multimodal_value(value: Any) -> Tuple[str, List[Any]]:
    if isinstance(value, dict):
        return value.get("text") or "", value.get("files") or []
    return "", []


def analyze_quiz(
    question_input: Dict[str, Any],
    figure_input: List[Any],
    solution_input: Dict[str, Any],
    backend_mode: str,
    api_region: str,
    api_base_url: str,
    api_model_name: str,
    api_key: str,
    save_api_key: bool,
    server_url: str,
    model_name: str,
    llama_server_exe: str,
    local_model_path: str,
    local_mmproj_path: str,
    llama_ctx_size: int,
    llama_threads: int,
    system_prompt: str,
    max_new_tokens: int,
    temperature: float,
    top_p: float,
    max_pdf_pages: int,
    max_image_side: int,
):
    """
    Extract stem/options from question_input, extract answer/remark from solution_input,
    and save only figure_input into the JSON image field.

    Important design choice:
    - question_input images/PDFs are OCR sources for extracting stem/options only.
    - figure_input images are the actual question figures to be stored in q["image"].
      They are not sent to the model by default, to avoid confusing OCR-source images
      with the image attached to the final question.
    """
    backend_mode = backend_mode or BACKEND_LOCAL
    empty_preview = ""
    yield f"正在准备后端：{backend_mode}...", "", "", "", "", "", "", empty_preview, {}, None

    if backend_mode == BACKEND_API:
        effective_server_url = resolve_api_base_url(api_region, api_base_url)
        effective_model_name = (api_model_name or DEFAULT_API_MODEL_NAME).strip()
        effective_api_key = (api_key or os.getenv("DASHSCOPE_API_KEY", "")).strip()
        if not effective_api_key:
            msg = "远程 API 模式需要 API Key。请在右侧 API key 输入框填写，或设置环境变量 DASHSCOPE_API_KEY。"
            yield "API Key 缺失：无法继续分析。", "", "", "", "", "", "", msg, {"error": "missing api key"}, None
            return
        try:
            cache_msg = save_backend_settings(
                backend_mode=backend_mode,
                api_region=api_region,
                api_base_url=api_base_url,
                api_model_name=api_model_name,
                api_key=api_key,
                save_api_key=bool(save_api_key),
                server_url=server_url,
                model_name=model_name,
                local_model_path=local_model_path,
                local_mmproj_path=local_mmproj_path,
            )
        except Exception as e:
            cache_msg = f"缓存设置失败：{type(e).__name__}: {e}"
        server_msg = (
            "远程 API 模式已就绪。\n"
            f"base_url={effective_server_url}\n"
            f"model={effective_model_name}\n"
            f"{cache_msg}"
        )
        yield "远程 API 已就绪。开始处理上传文件...", "", "", "", "", "", "", server_msg, {"backend": "api"}, None
    else:
        effective_server_url = server_url
        effective_model_name = (model_name or local_model_name_from_path(local_model_path)).strip()
        effective_api_key = ""
        ok, server_msg = ensure_llama_server_running(
            llama_server_exe=llama_server_exe,
            model_path=local_model_path,
            mmproj_path=local_mmproj_path,
            server_url=server_url,
            ctx_size=int(llama_ctx_size or DEFAULT_LLAMA_CONTEXT),
            threads=int(llama_threads or DEFAULT_LLAMA_THREADS),
            api_key="",
        )
        if not ok:
            yield (
                "本地 llama-server 未就绪：无法继续分析。",
                "", "", "", "", "", "",
                server_msg,
                {"error": "local llama-server not ready"},
                None,
            )
            return
        try:
            save_backend_settings(
                backend_mode=backend_mode,
                api_region=api_region,
                api_base_url=api_base_url,
                api_model_name=api_model_name,
                api_key=api_key,
                save_api_key=bool(save_api_key),
                server_url=server_url,
                model_name=effective_model_name,
                local_model_path=local_model_path,
                local_mmproj_path=local_mmproj_path,
            )
        except Exception:
            pass
        yield "本地 llama-server 已就绪。开始处理上传文件...", "", "", "", "", "", "", server_msg, {"backend": "local"}, None

    q_text, q_files = parse_multimodal_value(question_input)
    fig_text, fig_files = parse_multimodal_value(figure_input)
    s_text, s_files = parse_multimodal_value(solution_input)

    # OCR/source images: used only for extracting text.
    q_ocr_images, q_warnings = files_to_images(q_files, max_pdf_pages=max_pdf_pages, max_image_side=max_image_side)
    s_images, s_warnings = files_to_images(s_files, max_pdf_pages=max_pdf_pages, max_image_side=max_image_side)

    # Figure images: saved into JSON image field; not used as OCR source.
    # figure_input is a MultimodalTextbox so it supports upload/drag/paste in the browser.
    fig_images, fig_warnings = files_to_images(fig_files, max_pdf_pages=1, max_image_side=max_image_side)
    if fig_text.strip():
        fig_warnings.append("题目附图栏中的文字被忽略；该栏只用于保存题目附图。")
    image_rel = save_question_image_for_json(fig_images)

    status = (
        f"题干/选项 OCR 来源 {len(q_ocr_images)} 张，"
        f"题目附图 {len(fig_images)} 张，"
        f"答案/解析图片 {len(s_images)} 张。开始调用 {backend_mode} 提取题干和选项..."
    )
    yield status, "", "", "", "", "", image_rel, empty_preview, {}, None

    t_q0 = time.perf_counter()
    try:
        raw_q = run_vl_once(
            server_url=effective_server_url,
            model_name=effective_model_name,
            api_key=effective_api_key,
            prompt=build_question_prompt(q_text),
            images=q_ocr_images,
            system_prompt=system_prompt,
            max_new_tokens=max_new_tokens,
            temperature=temperature,
            top_p=top_p,
        )
    except Exception as e:
        err = (
            "[Question extraction ERROR]\n"
            f"{type(e).__name__}: {e}\n\n"
            "建议：本地模式请点击 Run diagnostics；API 模式请检查 base_url、API Key、模型名和图片输入支持。"
        )
        yield "调用后端失败：题干/选项提取未完成。", "", "", "", "", "", image_rel, err, {"error": str(e)}, None
        return
    q_elapsed = time.perf_counter() - t_q0
    if backend_mode == BACKEND_API:
        try:
            upsert_api_profile(
                api_region=api_region,
                api_base_url=api_base_url,
                api_model_name=api_model_name,
                api_key=effective_api_key,
                save_api_key=bool(save_api_key),
            )
        except Exception:
            pass
    q_obj = extract_json_object(raw_q)

    stem = str(q_obj.get("stem") or q_text or "").strip()
    options = normalize_options(q_obj.get("options", []))
    q_tags = normalize_tags(q_obj.get("category_tags", []))

    options_lines = "\n".join(options)
    status = "题干和选项提取完成。开始提取答案/解析并分类..."
    yield status, stem, options_lines, "", "，".join(q_tags), "", image_rel, raw_q, q_obj, None

    raw_s = ""
    s_obj: Dict[str, Any] = {}
    answer = ""
    explanation = ""
    s_tags: List[str] = []
    s_ocr = ""
    s_elapsed = 0.0

    if s_text.strip() or s_images:
        t_s0 = time.perf_counter()
        try:
            raw_s = run_vl_once(
                server_url=effective_server_url,
                model_name=effective_model_name,
                api_key=effective_api_key,
                prompt=build_solution_prompt(s_text, stem, options),
                images=s_images,
                system_prompt=system_prompt,
                max_new_tokens=max_new_tokens,
                temperature=temperature,
                top_p=top_p,
            )
        except Exception as e:
            raw_s = f"[Solution extraction ERROR]\n{type(e).__name__}: {e}"
            s_elapsed = time.perf_counter() - t_s0
            s_obj = {}
        else:
            s_elapsed = time.perf_counter() - t_s0
            s_obj = extract_json_object(raw_s)
        answer = str(s_obj.get("answer") or "").strip()
        explanation = str(s_obj.get("explanation") or s_text or "").strip()
        s_tags = normalize_tags(s_obj.get("category_tags", []))


    tags = []
    for t in q_tags + s_tags:
        if t and t not in tags:
            tags.append(t)
    if not tags:
        tags = ["其他"]

    remark = format_remark(tags=tags, explanation=explanation)

    final_q = {
        "stem": stem,
        "image": image_rel,
        "options": options,
        "answer": answer,
        "remark": remark,
    }

    raw_all = f"[Timing]\nquestion_call_seconds={q_elapsed:.2f}\nsolution_call_seconds={s_elapsed:.2f}\n\n[Question extraction raw output]\n" + raw_q
    if raw_s:
        raw_all += "\n\n[Solution extraction raw output]\n" + raw_s
    all_warnings = q_warnings + fig_warnings + s_warnings
    if all_warnings:
        raw_all += "\n\n[Warnings]\n" + "\n".join(all_warnings)

    preview = json.dumps(final_q, ensure_ascii=False, indent=2)
    status = f"完成。题干调用 {q_elapsed:.1f}s，解析调用 {s_elapsed:.1f}s。请检查并手动修正右侧字段，然后点击保存。"

    yield status, stem, options_lines, answer, "，".join(tags), remark, image_rel, raw_all, final_q, preview

def build_question_from_fields(stem: str, options_text: str, answer: str, category_text: str, remark: str, image_path: str) -> Dict[str, Any]:
    options = [x.strip() for x in options_text.splitlines() if x.strip()]
    # Ensure categories are present at the beginning of remark.
    cat_tags = normalize_tags(category_text)
    existing_remark = remark.strip()
    tag_line = " ".join(f"#{t}" for t in cat_tags if t)
    if tag_line and not existing_remark.startswith("#"):
        final_remark = (tag_line + "\n" + existing_remark).strip()
    elif tag_line:
        # avoid duplicating exact tags if already present
        final_remark = existing_remark
        for t in cat_tags:
            if f"#{t}" not in final_remark:
                final_remark = f"#{t} " + final_remark
    else:
        final_remark = existing_remark

    return {
        "stem": stem.strip(),
        "image": image_path.strip(),
        "options": options,
        "answer": answer.strip(),
        "remark": final_remark,
    }


def update_preview(stem: str, options_text: str, answer: str, category_text: str, remark: str, image_path: str):
    q = build_question_from_fields(stem, options_text, answer, category_text, remark, image_path)
    return json.dumps(q, ensure_ascii=False, indent=2), q


def save_current_question(stem: str, options_text: str, answer: str, category_text: str, remark: str, image_path: str):
    q = build_question_from_fields(stem, options_text, answer, category_text, remark, image_path)
    if not q["stem"]:
        return "保存失败：stem 为空。", json.dumps(q, ensure_ascii=False, indent=2), q
    idx = append_question(q)
    return f"已保存到 {DATA_FILE}，当前共 {idx} 道题。", json.dumps(q, ensure_ascii=False, indent=2), q


def clear_entry():
    return None, None, None, "", "", "", "", "", "", "", {}, ""


def make_question_label(global_idx: int, q: Dict[str, Any]) -> str:
    """Compact label used in CheckboxGroup. Prefix is 1-based original question id."""
    stem = str(q.get("stem", "")).replace("\n", " ").strip()
    answer = str(q.get("answer", "")).replace("\n", " ").strip()
    if len(stem) > 88:
        stem = stem[:88] + "..."
    return f"{global_idx + 1} | {stem} | 答案:{answer}"


def parse_question_label(label: str) -> Optional[int]:
    """Return 0-based original question index from a CheckboxGroup label."""
    try:
        first = str(label).split("|", 1)[0].strip()
        return int(first) - 1
    except Exception:
        return None


def question_rows_from_indices(indices: List[int]) -> List[List[Any]]:
    all_data = load_data()
    rows = []
    for pos, idx in enumerate(indices, 1):
        if 0 <= idx < len(all_data):
            q = all_data[idx]
            rows.append([
                pos,
                idx + 1,
                q.get("stem", "")[:100],
                " / ".join(q.get("options", []))[:140],
                q.get("answer", "")[:24],
                q.get("remark", "")[:140],
                q.get("image", ""),
            ])
    return rows


def remove_choices_from_indices(indices: List[int]) -> List[str]:
    all_data = load_data()
    choices = []
    for idx in indices:
        if 0 <= idx < len(all_data):
            choices.append(make_question_label(idx, all_data[idx]))
    return choices


def normalize_search_text(text: Any) -> str:
    """Normalize text for robust keyword matching."""
    text = str(text or "").lower()
    # Remove common whitespace noise while keeping Chinese characters intact.
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def split_search_query(keyword: str) -> Tuple[List[str], str]:
    """
    Parse user query into tokens.

    Default semantics: AND search.
      脊椎动物 骨骼  -> match records containing both tokens anywhere.

    If the query contains |, ／, /, or the word OR, use OR search.
      脊椎动物|骨骼 -> match either token.
    """
    keyword = normalize_search_text(keyword)
    if not keyword:
        return [], "recent"

    # Explicit OR syntax.
    if "|" in keyword or "/" in keyword or " or " in keyword:
        tokens = re.split(r"\s*(?:\||/|or)\s*", keyword)
        tokens = [t.strip() for t in tokens if t.strip()]
        return tokens, "OR"

    # Default: split by spaces, commas, Chinese comma, semicolon.
    tokens = re.split(r"[\s,，;；]+", keyword)
    tokens = [t.strip() for t in tokens if t.strip()]
    return tokens, "AND"


def question_search_corpus(q: Dict[str, Any]) -> str:
    fields = [
        q.get("stem", ""),
        q.get("answer", ""),
        q.get("remark", ""),
        q.get("image", ""),
        " ".join(str(x) for x in q.get("options", [])),
    ]
    return normalize_search_text(" ".join(fields))


def question_matches_tokens(q: Dict[str, Any], tokens: List[str], mode: str) -> bool:
    if not tokens:
        return True
    corpus = question_search_corpus(q)
    if mode == "OR":
        return any(t in corpus for t in tokens)
    return all(t in corpus for t in tokens)


def search_database(keyword: str, limit: int = 80):
    """
    Lightweight database browsing.

    Search syntax:
      - Empty query: show recent questions.
      - Multiple keywords separated by spaces/commas: AND search.
        Example: 脊椎动物 骨骼 -> both must appear in stem/options/answer/remark/image.
      - Use |, /, or OR for OR search.
        Example: 脊椎动物|骨骼 -> either can match.
    """
    all_data = load_data()
    total = len(all_data)
    keyword = (keyword or "").strip()
    limit = int(limit or 80)

    indexed_data = list(enumerate(all_data))
    tokens, mode_kind = split_search_query(keyword)

    if tokens:
        filtered = [
            (idx, q) for idx, q in indexed_data
            if question_matches_tokens(q, tokens, mode_kind)
        ]
        token_text = " + ".join(tokens) if mode_kind == "AND" else " | ".join(tokens)
        mode = f"搜索({mode_kind})：{token_text}"
    else:
        filtered = list(reversed(indexed_data))
        mode = "最近题目"

    shown = filtered[:limit]
    rows = []
    choices = []
    for display_i, (idx, q) in enumerate(shown, 1):
        rows.append([
            idx + 1,
            q.get("stem", "")[:90],
            " / ".join(q.get("options", []))[:120],
            q.get("answer", "")[:20],
            q.get("remark", "")[:120],
            q.get("image", ""),
        ])
        choices.append(make_question_label(idx, q))

    status = (
        f"{mode}：显示 {len(rows)} / 命中 {len(filtered)} / 总计 {total} 道题。"
        "多关键词默认取交集；如需并集，用 | 或 / 分隔。"
        "可点击 Select all visible results 一键勾选当前显示结果。"
    )
    return rows, status, gr.update(choices=choices, value=[]), choices


def select_all_search_results(current_result_choices: Optional[List[str]]):
    choices = list(current_result_choices or [])
    return gr.update(value=choices), f"已勾选当前显示的 {len(choices)} 道搜索结果；点击 Add checked to selected list 加入选中列表。"


def clear_search_checks():
    return gr.update(value=[]), "已清空当前搜索结果勾选。"


def add_checked_to_selection(checked_labels: List[str], selected_indices: Optional[List[int]]):
    selected_indices = list(selected_indices or [])
    checked_labels = checked_labels or []

    added = 0
    for label in checked_labels:
        idx = parse_question_label(label)
        if idx is not None and idx not in selected_indices:
            selected_indices.append(idx)
            added += 1

    selected_indices = [idx for idx in selected_indices if isinstance(idx, int)]
    rows = question_rows_from_indices(selected_indices)
    remove_choices = remove_choices_from_indices(selected_indices)
    status = f"已添加 {added} 道题；当前选中 {len(rows)} 道题。"
    return selected_indices, rows, gr.update(choices=remove_choices, value=[]), status


def remove_checked_from_selection(remove_labels: List[str], selected_indices: Optional[List[int]]):
    selected_indices = list(selected_indices or [])
    to_remove = set()
    for label in remove_labels or []:
        idx = parse_question_label(label)
        if idx is not None:
            to_remove.add(idx)

    selected_indices = [idx for idx in selected_indices if idx not in to_remove]
    rows = question_rows_from_indices(selected_indices)
    remove_choices = remove_choices_from_indices(selected_indices)
    status = f"已移除 {len(to_remove)} 道题；当前选中 {len(rows)} 道题。"
    return selected_indices, rows, gr.update(choices=remove_choices, value=[]), status


def clear_selection():
    return [], [], gr.update(choices=[], value=[]), "已清空选中列表。"


def database_initial_state():
    total = len(load_data())
    return [], f"题库共 {total} 道题。点击 Search / Reload 加载最近题目；为避免卡顿，不再自动渲染全部题库。", gr.update(choices=[], value=[])


def add_questions_to_doc(doc, questions: List[Dict[str, Any]]):
    for i, q in enumerate(questions, 1):
        doc.add_paragraph(f"{i}. {q.get('stem', '')}")
        img_rel = q.get("image", "")
        if img_rel:
            img_path = APP_DIR / img_rel
            if img_path.exists():
                try:
                    doc.add_picture(str(img_path), width=Inches(2.8))
                except Exception:
                    doc.add_paragraph("[图片无法显示]")
        for j, opt in enumerate(q.get("options", [])):
            doc.add_paragraph(f"  {chr(65 + j)}. {opt}")
        doc.add_paragraph("")

    doc.add_paragraph("\n答案与备注：")
    for i, q in enumerate(questions, 1):
        doc.add_paragraph(f"{i}. 答案: {q.get('answer', '无')}    备注: {q.get('remark', '无')}")


def export_selected_word(selected_indices: Optional[List[int]]):
    if Document is None or Inches is None:
        return "导出失败：缺少 python-docx。请执行 pip install python-docx。", None
    all_data = load_data()
    selected_indices = list(selected_indices or [])
    questions = [all_data[idx] for idx in selected_indices if 0 <= idx < len(all_data)]
    if not questions:
        return "没有选中任何题目。请先搜索、勾选并 Add checked to selected list。", None

    out_path = EXPORT_DIR / f"quiz_selected_export_{now_stamp()}.docx"
    doc = Document()
    add_questions_to_doc(doc, questions)
    doc.save(out_path)
    return f"已导出选中的 {len(questions)} 道题。", str(out_path)


def export_all_word():
    if Document is None or Inches is None:
        return "导出失败：缺少 python-docx。请执行 pip install python-docx。", None
    data = load_data()
    if not data:
        return "没有题目可导出。", None

    out_path = EXPORT_DIR / f"quiz_export_all_{now_stamp()}.docx"
    doc = Document()
    add_questions_to_doc(doc, data)
    doc.save(out_path)
    return f"已导出全部 {len(data)} 道题。", str(out_path)


def show_entry_view():
    """Switch view without loading database or model."""
    return gr.update(visible=True), gr.update(visible=False), "当前视图：AI 录入。切换动作不读取题库、不调用模型。"



def show_entry_view():
    """Switch view without loading database or model."""
    return (
        gr.update(visible=True),
        gr.update(visible=False),
        gr.update(visible=False),
        "当前视图：AI 录入。题库表格只在点击 Search / Reload 后加载。",
    )


def show_db_view():
    """Switch view without loading database or model."""
    return (
        gr.update(visible=False),
        gr.update(visible=True),
        gr.update(visible=False),
        "当前视图：题库浏览与导出。点击 Search / Reload 后才加载表格。",
    )


def show_ai_select_view():
    """Switch to AI-assisted selection view."""
    return (
        gr.update(visible=False),
        gr.update(visible=False),
        gr.update(visible=True),
        "当前视图：AI 辅助选题。AI 只拆解 AND/OR 关键词逻辑树和数量，固定程序负责检索与选定。",
    )


# =========================================================
# AI-assisted question selection
# =========================================================


def build_ai_selection_prompt(user_request: str, default_count: int, compact: bool = False) -> str:
    """Prompt the model to output a small boolean query tree instead of flat keyword lists."""
    default_count = int(default_count or 50)
    examples = """
示例1：用户：脊椎动物 骨骼 皮肤，50题
输出：{"count":50,"query":{"op":"AND","clauses":[{"name":"主题范围","op":"OR","keywords":["脊椎动物","脊索动物","鱼类","两栖","爬行","鸟类","哺乳"]},{"name":"考查模块","op":"OR","keywords":["骨骼","皮肤"]}]},"exclude_keywords":[],"notes":"脊椎动物是上位主题；骨骼和皮肤是同级模块，二者为 OR。"}
示例2：用户：同时涉及骨骼和皮肤的脊椎动物题目
输出：{"count":%d,"query":{"op":"AND","clauses":[{"name":"主题范围","op":"OR","keywords":["脊椎动物","脊索动物","鱼类","两栖","爬行","鸟类","哺乳"]},{"name":"骨骼","op":"OR","keywords":["骨骼"]},{"name":"皮肤","op":"OR","keywords":["皮肤"]}]},"exclude_keywords":[],"notes":"用户明确要求同时涉及骨骼和皮肤，因此两个模块用 AND。"}
示例3：用户：鱼类或两栖类的呼吸系统题
输出：{"count":%d,"query":{"op":"AND","clauses":[{"name":"动物类群","op":"OR","keywords":["鱼类","两栖"]},{"name":"考查模块","op":"OR","keywords":["呼吸"]}]},"exclude_keywords":[],"notes":"鱼类和两栖类是并列类群；呼吸系统是限定模块。"}
""" % (default_count, default_count)
    if compact:
        return f"""
只输出严格 JSON。不要解释。任务：把组卷需求转成布尔关键词树，程序会按树检索题库。
规则：上位主题 AND 子模块；同级候选词 OR；“同时/都涉及/并且”才用 AND；“或/或者/、/和”列举同级模块通常 OR；“不要/排除”放 exclude_keywords。
常用扩展：脊椎动物=>[脊椎动物,脊索动物,鱼类,两栖,爬行,鸟类,哺乳]。
JSON格式：{{"count":{default_count},"query":{{"op":"AND","clauses":[{{"name":"主题范围","op":"OR","keywords":[]}},{{"name":"考查模块","op":"OR","keywords":[]}}]}},"exclude_keywords":[],"notes":""}}
{examples}
用户需求：{user_request.strip() if user_request.strip() else '选择一套生物题'}
""".strip()
    return f"""
你是题库检索需求拆解助手。你不选题、不解题、不改写题目，只把用户需求转换为固定程序可执行的布尔关键词树。

输出必须是严格 JSON，不要 Markdown，不要额外解释。

核心原则：
1. count：用户明确题数就用用户题数，否则用默认题数 {default_count}。
2. query：用布尔树表达检索关系。节点格式为 {{"name":"...","op":"AND或OR","keywords":[...],"clauses":[...]}}。
3. 上位主题与下位模块通常是 AND。例如“脊椎动物 骨骼 皮肤”= 脊椎动物相关 AND (骨骼 OR 皮肤)。
4. 同级类群/同级模块通常是 OR。例如“骨骼 皮肤”是两个同级模块；“鱼类 两栖”是两个同级类群。
5. 用户明确说“同时涉及、都涉及、并且、既...又...”时，同级模块可以改成 AND。
6. 宽泛主题要扩展成题库可能出现的短词。例如“脊椎动物学”扩展为“脊椎动物、脊索动物、鱼类、两栖、爬行、鸟类、哺乳、比较解剖”。
7. exclude_keywords：只放用户明确排除的内容。
8. 关键词必须短，尽量是题干/remark 中可能出现的中文词，不要长句。
9. 不要输出 must_keywords / any_keywords；只输出 query 逻辑树。

JSON schema：
{{
  "count": 50,
  "query": {{
    "op": "AND",
    "clauses": [
      {{"name": "主题范围", "op": "OR", "keywords": []}},
      {{"name": "考查模块", "op": "OR", "keywords": []}}
    ]
  }},
  "exclude_keywords": [],
  "notes": "一句话说明 AND/OR 拆解逻辑"
}}

{examples}
用户需求：
{user_request.strip() if user_request.strip() else '选择一套生物题'}
""".strip()


def normalize_keyword_list(value: Any) -> List[str]:
    if isinstance(value, str):
        parts = re.split(r"[\s,，;；、/|]+", value)
    elif isinstance(value, list):
        parts = [str(x) for x in value]
    else:
        parts = []
    out: List[str] = []
    for x in parts:
        x = normalize_search_text(x).strip()
        if x and x not in out:
            out.append(x)
    return out


VERTEBRATE_TERMS = ["脊椎动物", "脊索动物", "鱼类", "两栖", "爬行", "鸟类", "哺乳", "比较解剖"]
MODULE_TERMS = [
    "骨骼", "皮肤", "循环", "呼吸", "排泄", "神经", "发育", "胚胎", "进化",
    "牙齿", "肌肉", "消化", "生殖", "免疫", "内分泌", "感觉", "运动", "器官系统",
]
TAXON_TERMS = ["鱼类", "两栖", "爬行", "鸟类", "哺乳", "哺乳类", "爬行类", "两栖类", "软骨鱼", "硬骨鱼"]


def make_query_leaf(name: str, op: str, keywords: List[str]) -> Dict[str, Any]:
    return {"name": name, "op": op.upper(), "keywords": normalize_keyword_list(keywords)}


def fallback_parse_selection_request(user_request: str, default_count: int) -> Dict[str, Any]:
    """Rule-based fallback when the small local model fails to return valid JSON."""
    text = user_request or ""
    m = re.search(r"(\d+)\s*[道题題个個]?", text)
    count = int(m.group(1)) if m else int(default_count or 50)

    exclude_keywords: List[str] = []
    for pat in [r"(?:不要|排除|不含|不包括)([^，,。；;]+)", r"(?:避免)([^，,。；;]+)"]:
        for hit in re.findall(pat, text):
            exclude_keywords.extend(normalize_keyword_list(hit))

    clauses: List[Dict[str, Any]] = []
    if any(t in text for t in ["脊椎", "脊索", "鱼", "两栖", "爬行", "鸟", "哺乳"]):
        topic_terms = VERTEBRATE_TERMS if any(t in text for t in ["脊椎", "脊索"]) else []
        taxa = [t for t in TAXON_TERMS if t in text]
        if taxa:
            topic_terms = taxa
        clauses.append(make_query_leaf("主题范围", "OR", topic_terms or VERTEBRATE_TERMS))

    module_hits = [t for t in MODULE_TERMS if t in text]
    if module_hits:
        if re.search(r"同时|都涉及|并且|既.*又", text) and len(module_hits) >= 2:
            for t in module_hits:
                clauses.append(make_query_leaf(t, "OR", [t]))
        else:
            clauses.append(make_query_leaf("考查模块", "OR", module_hits))

    if not clauses:
        tokens, _ = split_search_query(text)
        tokens = [t for t in tokens if not re.fullmatch(r"\d+", t)]
        clauses.append(make_query_leaf("用户关键词", "OR", tokens or [text.strip() or "生物"]))

    return {
        "count": count,
        "query": {"op": "AND", "clauses": clauses},
        "exclude_keywords": exclude_keywords,
        "notes": "AI 拆解失败或未返回合法 JSON，已使用规则兜底。",
    }


def sanitize_query_node(node: Any) -> Dict[str, Any]:
    """Normalize a model-produced query tree into a safe internal representation."""
    if not isinstance(node, dict):
        return {"op": "OR", "keywords": normalize_keyword_list(node)}
    op = str(node.get("op") or "OR").upper()
    if op not in {"AND", "OR"}:
        op = "OR"
    name = str(node.get("name") or "").strip()
    keywords = normalize_keyword_list(node.get("keywords", []))
    raw_clauses = node.get("clauses", [])
    clauses: List[Dict[str, Any]] = []
    if isinstance(raw_clauses, list):
        for child in raw_clauses:
            clean_child = sanitize_query_node(child)
            if clean_child.get("keywords") or clean_child.get("clauses"):
                clauses.append(clean_child)
    return {"name": name, "op": op, "keywords": keywords, "clauses": clauses}


def legacy_keywords_to_query(parsed: Dict[str, Any]) -> Dict[str, Any]:
    """Backward compatibility for older model output with must_keywords/any_keywords."""
    must_keywords = normalize_keyword_list(parsed.get("must_keywords", []))
    any_keywords = normalize_keyword_list(parsed.get("any_keywords", []))
    clauses: List[Dict[str, Any]] = []
    for kw in must_keywords:
        clauses.append(make_query_leaf(kw, "OR", [kw]))
    if any_keywords:
        clauses.append(make_query_leaf("扩展关键词", "OR", any_keywords))
    if not clauses:
        clauses.append(make_query_leaf("全部题目", "OR", []))
    return {"op": "AND", "clauses": clauses}


def ensure_query_plan(parsed: Dict[str, Any], user_request: str, default_count: int) -> Dict[str, Any]:
    """Return a valid {count, query, exclude_keywords, notes} plan."""
    if not isinstance(parsed, dict) or not parsed:
        return fallback_parse_selection_request(user_request, default_count)
    count = int(parsed.get("count") or default_count or 50)
    if isinstance(parsed.get("query"), dict):
        query = sanitize_query_node(parsed.get("query"))
    elif "must_keywords" in parsed or "any_keywords" in parsed:
        query = sanitize_query_node(legacy_keywords_to_query(parsed))
    else:
        return fallback_parse_selection_request(user_request, default_count)
    if not query.get("keywords") and not query.get("clauses"):
        return fallback_parse_selection_request(user_request, default_count)
    return {
        "count": count,
        "query": query,
        "exclude_keywords": normalize_keyword_list(parsed.get("exclude_keywords", [])),
        "notes": str(parsed.get("notes") or "").strip(),
    }


def evaluate_query_node(node: Dict[str, Any], corpus: str) -> bool:
    op = str(node.get("op") or "OR").upper()
    keywords = normalize_keyword_list(node.get("keywords", []))
    clauses = node.get("clauses", []) if isinstance(node.get("clauses", []), list) else []
    results: List[bool] = []
    if keywords:
        if op == "AND":
            results.append(all(k in corpus for k in keywords))
        else:
            results.append(any(k in corpus for k in keywords))
    for child in clauses:
        if isinstance(child, dict):
            results.append(evaluate_query_node(child, corpus))
    if not results:
        return True
    return all(results) if op == "AND" else any(results)


def flatten_query_keywords(node: Dict[str, Any], depth: int = 0) -> List[Tuple[str, int]]:
    out: List[Tuple[str, int]] = []
    for kw in normalize_keyword_list(node.get("keywords", [])):
        out.append((kw, depth))
    clauses = node.get("clauses", []) if isinstance(node.get("clauses", []), list) else []
    for child in clauses:
        if isinstance(child, dict):
            out.extend(flatten_query_keywords(child, depth + 1))
    return out


def query_tree_to_text(node: Dict[str, Any]) -> str:
    op = str(node.get("op") or "OR").upper()
    name = str(node.get("name") or "").strip()
    keywords = normalize_keyword_list(node.get("keywords", []))
    parts: List[str] = []
    if keywords:
        joiner = f" {op} "
        parts.append("(" + joiner.join(keywords) + ")" if len(keywords) > 1 else keywords[0])
    clauses = node.get("clauses", []) if isinstance(node.get("clauses", []), list) else []
    for child in clauses:
        if isinstance(child, dict):
            parts.append(query_tree_to_text(child))
    if not parts:
        expr = "ALL"
    elif len(parts) == 1:
        expr = parts[0]
    else:
        expr = "(" + f" {op} ".join(parts) + ")"
    return f"{name}:{expr}" if name else expr


def select_question_indices_by_query_tree(
    query: Dict[str, Any],
    exclude_keywords: List[str],
    count: int,
) -> Tuple[List[int], List[List[Any]], str]:
    all_data = load_data()
    count = max(1, int(count or 50))
    query = sanitize_query_node(query)
    exclude_keywords = normalize_keyword_list(exclude_keywords)
    flat_keywords = flatten_query_keywords(query)

    scored: List[Tuple[int, int, Dict[str, Any]]] = []
    for idx, q in enumerate(all_data):
        corpus = question_search_corpus(q)
        if exclude_keywords and any(k in corpus for k in exclude_keywords):
            continue
        if not evaluate_query_node(query, corpus):
            continue
        score = 0
        for kw, depth in flat_keywords:
            if kw in corpus:
                score += max(8, 40 - depth * 6)
        if "脊椎动物学" in corpus:
            score += 30
        if "文献题" in corpus or "图表题" in corpus:
            score += 5
        scored.append((score, idx, q))

    scored.sort(key=lambda x: (-x[0], x[1]))
    indices = [idx for _, idx, _ in scored[:count]]
    rows = question_rows_from_indices(indices)
    status = (
        f"逻辑树检索完成：选中 {len(indices)} / 目标 {count} / 题库总计 {len(all_data)}。\n"
        f"query={query_tree_to_text(query)}\n"
        f"exclude={exclude_keywords or '无'}"
    )
    return indices, rows, status


# Backward-compatible wrapper. Kept in case other UI callbacks still call the old function name.
def select_question_indices_by_keywords(
    must_keywords: List[str],
    any_keywords: List[str],
    exclude_keywords: List[str],
    count: int,
) -> Tuple[List[int], List[List[Any]], str]:
    clauses: List[Dict[str, Any]] = []
    for kw in normalize_keyword_list(must_keywords):
        clauses.append(make_query_leaf(kw, "OR", [kw]))
    any_keywords = normalize_keyword_list(any_keywords)
    if any_keywords:
        clauses.append(make_query_leaf("扩展关键词", "OR", any_keywords))
    query = {"op": "AND", "clauses": clauses or [make_query_leaf("全部题目", "OR", [])]}
    return select_question_indices_by_query_tree(query, exclude_keywords, count)


def ai_plan_and_select_questions(
    user_request: str,
    default_count: int,
    backend_mode: str,
    api_region: str,
    api_base_url: str,
    api_model_name: str,
    api_key: str,
    server_url: str,
    model_name: str,
    llama_server_exe: str,
    local_model_path: str,
    local_mmproj_path: str,
    llama_ctx_size: int,
    llama_threads: int,
    system_prompt: str,
    temperature: float,
    top_p: float,
) -> Tuple[str, Dict[str, Any], str, List[List[Any]], List[int], Any, List[str]]:
    """Use AI only to decompose the selection request; deterministic code performs retrieval.

    Returns current AI candidates plus CheckboxGroup choices so the user can
    discard items before merging selected candidates into a persistent total pool.
    """
    user_request = (user_request or "").strip()
    if not user_request:
        return "请先输入选题需求。", {}, "", [], [], gr.update(choices=[], value=[]), []

    backend_mode = backend_mode or BACKEND_LOCAL
    raw = ""
    parsed: Dict[str, Any] = {}
    try:
        if backend_mode == BACKEND_API:
            effective_server_url = resolve_api_base_url(api_region, api_base_url)
            effective_model_name = (api_model_name or DEFAULT_API_MODEL_NAME).strip()
            effective_api_key = (api_key or os.getenv("DASHSCOPE_API_KEY", "")).strip()
            if not effective_api_key:
                raise RuntimeError("远程 API 模式需要 API key。")
            compact_prompt = False
        else:
            effective_server_url = server_url
            effective_model_name = (model_name or local_model_name_from_path(local_model_path)).strip()
            effective_api_key = ""
            compact_prompt = True
            ok, msg = ensure_llama_server_running(
                llama_server_exe=llama_server_exe,
                model_path=local_model_path,
                mmproj_path=local_mmproj_path,
                server_url=server_url,
                ctx_size=int(llama_ctx_size or DEFAULT_LLAMA_CONTEXT),
                threads=int(llama_threads or DEFAULT_LLAMA_THREADS),
                api_key="",
            )
            if not ok:
                raise RuntimeError(msg)

        raw = run_vl_once(
            server_url=effective_server_url,
            model_name=effective_model_name,
            api_key=effective_api_key,
            prompt=build_ai_selection_prompt(user_request, int(default_count or 50), compact=compact_prompt),
            images=[],
            system_prompt="只输出严格 JSON，不要输出解释或 Markdown。",
            max_new_tokens=640 if compact_prompt else 1024,
            temperature=0.0,
            top_p=1.0,
        )
        parsed = extract_json_object(raw)
        if not parsed:
            parsed = fallback_parse_selection_request(user_request, int(default_count or 50))
            raw = raw + "\n\n[Fallback]\n模型未返回合法 JSON，已使用规则兜底。"
    except Exception as e:
        parsed = fallback_parse_selection_request(user_request, int(default_count or 50))
        raw = f"[AI planning failed]\n{type(e).__name__}: {e}\n\n[Fallback]\n已使用规则兜底。"

    plan = ensure_query_plan(parsed, user_request, int(default_count or 50))
    count = int(plan.get("count") or default_count or 50)
    query = sanitize_query_node(plan.get("query", {}))
    exclude_keywords = normalize_keyword_list(plan.get("exclude_keywords", []))

    indices, rows, select_status = select_question_indices_by_query_tree(
        query=query,
        exclude_keywords=exclude_keywords,
        count=count,
    )
    plan_summary = json.dumps({
        "count": count,
        "query": query,
        "query_text": query_tree_to_text(query),
        "exclude_keywords": exclude_keywords,
        "notes": plan.get("notes", ""),
    }, ensure_ascii=False, indent=2)
    current_choices = remove_choices_from_indices(indices)
    status = (
        "AI 已完成 AND/OR 逻辑树拆解，固定程序已按逻辑树检索候选题。\n"
        + select_status
        + "\n请在 AI current candidates 中取消不想保留的题目，再点击 Add checked to total pool。"
    )
    return (
        status,
        plan,
        raw + "\n\n[Query plan]\n" + plan_summary,
        rows,
        indices,
        gr.update(choices=current_choices, value=current_choices),
        current_choices,
    )


def discard_current_ai_candidates():
    """Clear the current AI candidate batch without touching the persistent total pool."""
    return (
        [],
        [],
        gr.update(choices=[], value=[]),
        [],
        "已丢弃当前 AI 候选批次；总备选池未改变。",
    )


def add_current_ai_checked_to_pool(checked_labels: List[str], pool_indices: Optional[List[int]]):
    """Merge checked AI candidates into the persistent total pool."""
    return add_checked_to_selection(checked_labels, pool_indices)


def remove_checked_from_ai_pool(remove_labels: List[str], pool_indices: Optional[List[int]]):
    """Remove checked questions from the persistent AI total pool."""
    return remove_checked_from_selection(remove_labels, pool_indices)


def clear_ai_pool():
    """Clear persistent AI total pool."""
    return clear_selection()


# =========================================================
# UI
# =========================================================

# This version intentionally avoids gr.Tab.
# It uses a front-end JS visibility toggle, so switching between entry and database views
# does not call Python and should feel immediate.

NAV_HTML = """
<div id="topbar">
  <h1>Qwen Based Quiz Builder</h1>
  <p>题库录入工作台：题干/选项 OCR、题目附图、答案解析抽取、题库浏览与 Word 导出。</p>
</div>
"""

CUSTOM_CSS = CUSTOM_CSS + """
#view-nav {
  padding: 10px 12px;
  margin-bottom: 14px;
  border: 1px solid #e6e8ef;
  border-radius: 18px;
  background: #ffffff;
}
#entry-view, #db-view, #ai-select-view { width: 100%; }
#db-panel {
  border: 1px solid #e6e8ef;
  border-radius: 22px;
  padding: 16px;
  background: #ffffff;
  min-height: 820px;
}
#db-controls {
  padding: 12px;
  border-radius: 18px;
  background: #f8f9fc;
  border: 1px solid #e6e8ef;
  margin-bottom: 12px;
}
#db-table { max-height: 620px; overflow: auto; }
"""

with gr.Blocks(title="Local Qwen-VL Quiz Builder - llama-server checked") as demo:
    extracted_state = gr.State({})

    gr.HTML(NAV_HTML)

    with gr.Row(elem_id="view-nav"):
        nav_entry_btn = gr.Button("AI 录入", variant="primary", scale=1)
        nav_db_btn = gr.Button("题库浏览与导出", variant="secondary", scale=1)
        nav_ai_select_btn = gr.Button("AI 辅助选题", variant="secondary", scale=1)
        nav_status = gr.Textbox(
            value="当前视图：AI 录入。题库表格只在点击 Search / Reload 后加载。",
            interactive=False,
            show_label=False,
            scale=5,
        )

    with gr.Column(elem_id="entry-view", visible=True) as entry_view:
        with gr.Row():
            with gr.Column(scale=4, elem_id="left-panel", min_width=420):
                gr.Markdown("## 1. 题干与选项 OCR 来源")
                question_input = gr.MultimodalTextbox(
                    label="题干/选项输入：文字、截图或 PDF；只用于提取 stem/options，不会保存到 JSON 的 image 字段",
                    placeholder="例如：请从这张截图中提取题干和选项。也可以直接粘贴题目文字。",
                    file_count="multiple",
                    file_types=["image", ".pdf"],
                    elem_id="q-input",
                )
                gr.Markdown("## 2. 题目附图")
                figure_input = gr.MultimodalTextbox(
                    label="题目真正附带的图：可上传、拖入或粘贴；保存到 questions.json 的 image 字段；默认不参与题干 OCR",
                    placeholder="在这里粘贴/拖入/上传题目真正附带的图。这里的图片会保存到 JSON 的 image 字段。",
                    file_count="multiple",
                    file_types=["image"],
                )
                gr.Markdown("## 3. 答案与解析")
                solution_input = gr.MultimodalTextbox(
                    label="答案/解析输入：可输入文字，也可上传解析图片或 PDF；没有则留空",
                    placeholder="例如：答案为 TTFT。解析：……",
                    file_count="multiple",
                    file_types=["image", ".pdf"],
                    elem_id="a-input",
                )
                with gr.Row():
                    analyze_btn = gr.Button("Analyze with llama-server", variant="primary")
                    clear_btn = gr.Button("Clear")
                status = gr.Textbox(label="Status", value="等待输入。", interactive=False)

                gr.Markdown("## JSON Preview")
                json_preview = gr.Code(
                    label="JSON preview",
                    language="json",
                    lines=16,
                    elem_id="json-preview",
                )

            with gr.Column(scale=5, elem_id="center-panel", min_width=560):
                gr.Markdown("## 4. 审校与保存")
                stem_box = gr.Textbox(label="stem / 题干", lines=6)
                options_box = gr.Textbox(label="options / 选项：每行一个选项，不要写 A/B/C/D", lines=7)
                with gr.Row():
                    answer_box = gr.Textbox(label="answer / 答案", scale=1, placeholder="如 TFFF 或 AC")
                    category_box = gr.Textbox(label="category tags / 分类", scale=2, placeholder="如 细胞生物学，信号转导")
                image_path_box = gr.Textbox(label="image / 题目图片保存路径", placeholder="images/xxx.png")
                remark_box = gr.Textbox(label="remark / 备注、分类、解析", lines=12)
                with gr.Row():
                    preview_btn = gr.Button("Update Preview")
                    save_btn = gr.Button("Save to questions.json", variant="primary")
                save_status = gr.Textbox(label="Save status", interactive=False)

            with gr.Column(scale=4, elem_id="right-panel", min_width=420):
                gr.Markdown("## 5. 后端、模型与输出")
                backend_mode = gr.Dropdown(
                    label="推理后端",
                    choices=[BACKEND_LOCAL, BACKEND_API],
                    value=config_get(BACKEND_CONFIG, "backend_mode", BACKEND_LOCAL),
                    interactive=True,
                )
                backend_hint = gr.Textbox(
                    label="后端状态提示",
                    value=("当前后端：远程 API。不会自动启动本地 llama-server。" if config_get(BACKEND_CONFIG, "backend_mode", BACKEND_LOCAL) == BACKEND_API else "当前后端：本地 llama-server。Analyze 时会自动检测/启动本地模型服务。"),
                    interactive=False,
                    lines=2,
                )

                with gr.Group(visible=(config_get(BACKEND_CONFIG, "backend_mode", BACKEND_LOCAL) != BACKEND_API)) as local_backend_group:
                    gr.Markdown("### 本地 llama-server / GGUF")
                    server_url = gr.Textbox(label="llama-server URL", value=config_get(BACKEND_CONFIG, "server_url", DEFAULT_LLAMA_SERVER_URL), lines=1)
                    model_name = gr.Textbox(label="本地模型名；通常可用 GGUF 文件名", value=config_get(BACKEND_CONFIG, "model_name", DEFAULT_LLAMA_MODEL_NAME), lines=1)

                    local_model_dropdown = gr.Dropdown(
                        label="选择 ./models 中的 GGUF 主模型",
                        choices=LOCAL_MODEL_CHOICES,
                        value=config_get(BACKEND_CONFIG, "local_model_path", LOCAL_MODEL_CHOICES[0] if LOCAL_MODEL_CHOICES else DEFAULT_LOCAL_GGUF_MODEL),
                        interactive=True,
                        allow_custom_value=True,
                    )
                    local_model_path = gr.Textbox(
                        label="本地 GGUF 主模型路径",
                        value=config_get(BACKEND_CONFIG, "local_model_path", LOCAL_MODEL_CHOICES[0] if LOCAL_MODEL_CHOICES else DEFAULT_LOCAL_GGUF_MODEL),
                        lines=2,
                    )
                    local_mmproj_dropdown = gr.Dropdown(
                        label="选择 ./models 中的 mmproj",
                        choices=LOCAL_MMPROJ_CHOICES,
                        value=config_get(BACKEND_CONFIG, "local_mmproj_path", LOCAL_MMPROJ_CHOICES[0] if LOCAL_MMPROJ_CHOICES else DEFAULT_LOCAL_MMPROJ),
                        interactive=True,
                        allow_custom_value=True,
                    )
                    local_mmproj_path = gr.Textbox(
                        label="mmproj 路径；图像输入必需，建议 mmproj-F16.gguf",
                        value=config_get(BACKEND_CONFIG, "local_mmproj_path", LOCAL_MMPROJ_CHOICES[0] if LOCAL_MMPROJ_CHOICES else DEFAULT_LOCAL_MMPROJ),
                        lines=2,
                    )
                    llama_server_exe = gr.Textbox(
                        label="llama-server.exe 路径；若已加入 PATH 可保持 llama-server.exe",
                        value=DEFAULT_LLAMA_SERVER_EXE,
                        lines=1,
                    )
                    with gr.Row():
                        refresh_models_btn = gr.Button("Refresh local models")
                        start_llama_btn = gr.Button("Start local llama-server", variant="secondary")
                    with gr.Row():
                        diag_llama_btn = gr.Button("Run diagnostics")
                        preview_cmd_btn = gr.Button("Preview command")
                    with gr.Row():
                        llama_ctx_size = gr.Slider(
                            label="llama ctx-size",
                            minimum=2048,
                            maximum=16384,
                            value=DEFAULT_LLAMA_CONTEXT,
                            step=1024,
                        )
                        llama_threads = gr.Slider(
                            label="CPU threads",
                            minimum=1,
                            maximum=max(1, os.cpu_count() or 16),
                            value=DEFAULT_LLAMA_THREADS,
                            step=1,
                        )
                    llama_start_status = gr.Textbox(label="local status / command / diagnostics", lines=10, interactive=False)

                with gr.Group(visible=(config_get(BACKEND_CONFIG, "backend_mode", BACKEND_LOCAL) == BACKEND_API)) as api_backend_group:
                    gr.Markdown("### 远程 OpenAI-compatible API")
                    api_profile_dropdown = gr.Dropdown(
                        label="已保存 API profile；选择后点击 Load",
                        choices=API_PROFILE_CHOICES,
                        value=None,
                        interactive=True,
                    )
                    with gr.Row():
                        load_api_profile_btn = gr.Button("Load selected API profile")
                        delete_api_profile_btn = gr.Button("Delete selected API profile")
                    api_region = gr.Dropdown(
                        label="API region endpoint",
                        choices=list(REGION_BASE_URLS.keys()),
                        value=config_get(BACKEND_CONFIG, "api_region", DEFAULT_API_REGION),
                        interactive=True,
                    )
                    api_base_url = gr.Textbox(
                        label="Custom base_url，可留空；留空则使用 region endpoint",
                        value=config_get(BACKEND_CONFIG, "api_base_url", ""),
                        placeholder="例如 https://dashscope.aliyuncs.com/compatible-mode/v1",
                    )
                    api_model_name = gr.Textbox(
                        label="API model name",
                        value=config_get(BACKEND_CONFIG, "api_model_name", DEFAULT_API_MODEL_NAME),
                    )
                    api_key = gr.Textbox(
                        label="API key；会按下方选项保存到本地缓存",
                        value=config_get(BACKEND_CONFIG, "api_key", os.getenv("DASHSCOPE_API_KEY", "")),
                        type="password",
                    )
                    save_api_key = gr.Checkbox(
                        label="Save API key to local cache",
                        value=bool(config_get(BACKEND_CONFIG, "save_api_key", True)),
                    )
                    with gr.Row():
                        save_backend_btn = gr.Button("Save backend/API settings")
                        test_save_api_profile_btn = gr.Button("Test API and save URL-key profile", variant="primary")
                    api_cache_status = gr.Textbox(label="API/cache status", lines=4, interactive=False)

                system_prompt = gr.Textbox(label="System prompt", value=DEFAULT_SYSTEM_PROMPT, lines=8)
                with gr.Row():
                    temperature = gr.Slider(label="temperature", minimum=0.0, maximum=1.0, value=0.0, step=0.05)
                    top_p = gr.Slider(label="top_p", minimum=0.1, maximum=1.0, value=0.9, step=0.05)
                max_new_tokens = gr.Slider(label="max_new_tokens", minimum=512, maximum=4096, value=1536, step=128)
                max_pdf_pages = gr.Slider(label="Max PDF pages", minimum=1, maximum=10, value=2, step=1)
                max_image_side = gr.Slider(label="Max image side", minimum=900, maximum=2400, value=1500, step=100)
                raw_output = gr.Textbox(label="Raw model output", lines=14)
                parsed_json = gr.JSON(label="Parsed model JSON")

    with gr.Column(elem_id="db-view", visible=False) as db_view:
        selected_indices_state = gr.State([])
        search_result_choices_state = gr.State([])
        with gr.Column(elem_id="db-panel"):
            gr.Markdown("## 题库浏览、勾选与导出")
            gr.Markdown("先搜索或加载最近题目，在搜索结果中勾选题目并加入 **Selected list**，最后只导出选中的题目。多关键词默认取交集，例如 `脊椎动物 骨骼`；并集搜索可写 `脊椎动物|骨骼`。")
            with gr.Row(elem_id="db-controls"):
                keyword_box = gr.Textbox(label="Search keyword", placeholder="多关键词搜索：如 脊椎动物 骨骼；并集：脊椎动物|骨骼", scale=4)
                db_limit = gr.Slider(label="Max rows", minimum=20, maximum=300, value=80, step=20, scale=2)
                search_btn = gr.Button("Search / Reload", variant="primary", scale=1)
                export_all_btn = gr.Button("Export all", scale=1)
            db_status = gr.Textbox(
                label="Database status",
                value="未加载表格。点击 Search / Reload 加载最近题目。",
                interactive=False,
            )
            db_table = gr.Dataframe(
                headers=["id", "stem", "options", "answer", "remark", "image"],
                datatype=["number", "str", "str", "str", "str", "str"],
                interactive=False,
                wrap=False,
                elem_id="db-table",
            )
            result_checks = gr.CheckboxGroup(
                label="Search results: check questions to add",
                choices=[],
                value=[],
            )
            with gr.Row():
                select_all_btn = gr.Button("Select all visible results", scale=1)
                clear_result_checks_btn = gr.Button("Clear result checks", scale=1)
                add_selected_btn = gr.Button("Add checked to selected list", variant="primary", scale=2)
                clear_selected_btn = gr.Button("Clear selected list", scale=1)

            gr.Markdown("### Selected list")
            selected_status = gr.Textbox(
                label="Selected status",
                value="当前未选中题目。",
                interactive=False,
            )
            selected_table = gr.Dataframe(
                headers=["selected #", "id", "stem", "options", "answer", "remark", "image"],
                datatype=["number", "number", "str", "str", "str", "str", "str"],
                interactive=False,
                wrap=False,
                elem_id="selected-table",
            )
            remove_checks = gr.CheckboxGroup(
                label="Selected list: check questions to remove",
                choices=[],
                value=[],
            )
            with gr.Row():
                remove_selected_btn = gr.Button("Remove checked from selected list", scale=2)
                export_selected_btn = gr.Button("Export selected to Word", variant="primary", scale=2)

            with gr.Row():
                export_status = gr.Textbox(label="Export status", interactive=False, scale=3)
                export_file = gr.File(label="Exported Word file", scale=2)

    with gr.Column(elem_id="ai-select-view", visible=False) as ai_select_view:
        # Current batch = the latest AI retrieval result. Total pool = manually curated across multiple AI runs.
        ai_current_indices_state = gr.State([])
        ai_current_choices_state = gr.State([])
        ai_pool_indices_state = gr.State([])
        with gr.Column(elem_id="db-panel"):
            gr.Markdown("## AI 辅助选题")
            gr.Markdown(
                "输入自然语言需求，例如“选择脊椎动物学的题目，50题”或“脊椎动物 骨骼相关题目”。"
                "AI 只负责拆解为 count + AND/OR 关键词逻辑树；随后由固定程序检索候选题。"
                "你可以人工取消不想保留的候选题，再加入总备选池；可多次 AI 选题、多轮人工筛选，最后导出总备选。"
            )
            with gr.Row(elem_id="db-controls"):
                ai_select_request = gr.Textbox(
                    label="选题需求",
                    value="选择脊椎动物学的题目，50题",
                    lines=3,
                    scale=5,
                )
                ai_default_count = gr.Slider(
                    label="默认题数",
                    minimum=5,
                    maximum=150,
                    value=50,
                    step=5,
                    scale=2,
                )
                ai_select_btn = gr.Button("AI 拆解需求并生成候选", variant="primary", scale=2)
            ai_select_status = gr.Textbox(label="AI selection status", lines=6, interactive=False)
            with gr.Row():
                ai_keyword_plan = gr.JSON(label="Parsed keyword plan")
                ai_raw_plan = gr.Textbox(label="Raw AI planning output", lines=12)

            gr.Markdown("### Current AI candidates：本轮 AI 候选")
            ai_current_table = gr.Dataframe(
                headers=["selected #", "id", "stem", "options", "answer", "remark", "image"],
                datatype=["number", "number", "str", "str", "str", "str", "str"],
                interactive=False,
                wrap=False,
                elem_id="ai-current-table",
            )
            ai_current_checks = gr.CheckboxGroup(
                label="AI current candidates: uncheck questions to discard before adding to total pool",
                choices=[],
                value=[],
            )
            with gr.Row():
                ai_select_all_current_btn = gr.Button("Select all current candidates", scale=1)
                ai_clear_current_checks_btn = gr.Button("Clear current checks", scale=1)
                ai_add_checked_to_pool_btn = gr.Button("Add checked to total pool", variant="primary", scale=2)
                ai_discard_current_btn = gr.Button("Discard current batch", scale=1)

            gr.Markdown("### Total pool：总备选池")
            ai_pool_status = gr.Textbox(
                label="Total pool status",
                value="总备选池为空。可以多次 AI 选题，将人工保留的题目加入这里，最后统一导出。",
                lines=2,
                interactive=False,
            )
            ai_pool_table = gr.Dataframe(
                headers=["selected #", "id", "stem", "options", "answer", "remark", "image"],
                datatype=["number", "number", "str", "str", "str", "str", "str"],
                interactive=False,
                wrap=False,
                elem_id="ai-pool-table",
            )
            ai_pool_remove_checks = gr.CheckboxGroup(
                label="Total pool: check questions to remove",
                choices=[],
                value=[],
            )
            with gr.Row():
                ai_remove_from_pool_btn = gr.Button("Remove checked from total pool", scale=2)
                ai_clear_pool_btn = gr.Button("Clear total pool", scale=1)
                ai_export_pool_btn = gr.Button("Export total pool to Word", variant="primary", scale=2)
            with gr.Row():
                ai_export_status = gr.Textbox(label="AI export status", interactive=False, scale=3)
                ai_export_file = gr.File(label="Exported Word file", scale=2)


    nav_entry_btn.click(
        fn=show_entry_view,
        inputs=[],
        outputs=[entry_view, db_view, ai_select_view, nav_status],
    )

    nav_db_btn.click(
        fn=show_db_view,
        inputs=[],
        outputs=[entry_view, db_view, ai_select_view, nav_status],
    )

    nav_ai_select_btn.click(
        fn=show_ai_select_view,
        inputs=[],
        outputs=[entry_view, db_view, ai_select_view, nav_status],
    )

    backend_mode.change(
        fn=update_backend_visibility,
        inputs=[backend_mode],
        outputs=[local_backend_group, api_backend_group, backend_hint],
    )

    load_api_profile_btn.click(
        fn=load_api_profile,
        inputs=[api_profile_dropdown],
        outputs=[api_region, api_base_url, api_model_name, api_key, save_api_key, api_cache_status],
    )

    delete_api_profile_btn.click(
        fn=delete_api_profile,
        inputs=[api_profile_dropdown],
        outputs=[api_profile_dropdown, api_cache_status],
    )

    test_save_api_profile_btn.click(
        fn=test_api_and_save_profile,
        inputs=[api_region, api_base_url, api_model_name, api_key, save_api_key],
        outputs=[api_profile_dropdown, api_cache_status],
    )

    refresh_models_btn.click(
        fn=refresh_local_model_choices,
        inputs=[],
        outputs=[local_model_dropdown, local_mmproj_dropdown, llama_start_status],
    )

    local_model_dropdown.change(
        fn=apply_local_model_choice,
        inputs=[local_model_dropdown, local_mmproj_path],
        outputs=[local_model_path, model_name, local_mmproj_path],
    )

    local_mmproj_dropdown.change(
        fn=lambda x: x or DEFAULT_LOCAL_MMPROJ,
        inputs=[local_mmproj_dropdown],
        outputs=[local_mmproj_path],
    )

    save_backend_btn.click(
        fn=save_backend_settings,
        inputs=[
            backend_mode,
            api_region,
            api_base_url,
            api_model_name,
            api_key,
            save_api_key,
            server_url,
            model_name,
            local_model_path,
            local_mmproj_path,
        ],
        outputs=[api_cache_status],
    )

    start_llama_btn.click(
        fn=start_local_llama_server,
        inputs=[
            llama_server_exe,
            local_model_path,
            local_mmproj_path,
            server_url,
            llama_ctx_size,
            llama_threads,
        ],
        outputs=[llama_start_status],
    )

    diag_llama_btn.click(
        fn=run_llama_diagnostics,
        inputs=[
            llama_server_exe,
            local_model_path,
            local_mmproj_path,
            server_url,
            llama_ctx_size,
            llama_threads,
            api_key,
        ],
        outputs=[llama_start_status],
    )

    preview_cmd_btn.click(
        fn=preview_llama_server_command,
        inputs=[
            llama_server_exe,
            local_model_path,
            local_mmproj_path,
            server_url,
            llama_ctx_size,
            llama_threads,
        ],
        outputs=[llama_start_status],
    )

    analyze_btn.click(
        fn=analyze_quiz,
        inputs=[
            question_input,
            figure_input,
            solution_input,
            backend_mode,
            api_region,
            api_base_url,
            api_model_name,
            api_key,
            save_api_key,
            server_url,
            model_name,
            llama_server_exe,
            local_model_path,
            local_mmproj_path,
            llama_ctx_size,
            llama_threads,
            system_prompt,
            max_new_tokens,
            temperature,
            top_p,
            max_pdf_pages,
            max_image_side,
        ],
        outputs=[
            status,
            stem_box,
            options_box,
            answer_box,
            category_box,
            remark_box,
            image_path_box,
            raw_output,
            parsed_json,
            json_preview,
        ],
    )

    ai_select_btn.click(
        fn=ai_plan_and_select_questions,
        inputs=[
            ai_select_request,
            ai_default_count,
            backend_mode,
            api_region,
            api_base_url,
            api_model_name,
            api_key,
            server_url,
            model_name,
            llama_server_exe,
            local_model_path,
            local_mmproj_path,
            llama_ctx_size,
            llama_threads,
            system_prompt,
            temperature,
            top_p,
        ],
        outputs=[
            ai_select_status,
            ai_keyword_plan,
            ai_raw_plan,
            ai_current_table,
            ai_current_indices_state,
            ai_current_checks,
            ai_current_choices_state,
        ],
    )

    ai_select_all_current_btn.click(
        fn=select_all_search_results,
        inputs=[ai_current_choices_state],
        outputs=[ai_current_checks, ai_select_status],
    )

    ai_clear_current_checks_btn.click(
        fn=clear_search_checks,
        inputs=[],
        outputs=[ai_current_checks, ai_select_status],
    )

    ai_discard_current_btn.click(
        fn=discard_current_ai_candidates,
        inputs=[],
        outputs=[
            ai_current_indices_state,
            ai_current_table,
            ai_current_checks,
            ai_current_choices_state,
            ai_select_status,
        ],
    )

    ai_add_checked_to_pool_btn.click(
        fn=add_current_ai_checked_to_pool,
        inputs=[ai_current_checks, ai_pool_indices_state],
        outputs=[ai_pool_indices_state, ai_pool_table, ai_pool_remove_checks, ai_pool_status],
    )

    ai_remove_from_pool_btn.click(
        fn=remove_checked_from_ai_pool,
        inputs=[ai_pool_remove_checks, ai_pool_indices_state],
        outputs=[ai_pool_indices_state, ai_pool_table, ai_pool_remove_checks, ai_pool_status],
    )

    ai_clear_pool_btn.click(
        fn=clear_ai_pool,
        inputs=[],
        outputs=[ai_pool_indices_state, ai_pool_table, ai_pool_remove_checks, ai_pool_status],
    )

    ai_export_pool_btn.click(
        fn=export_selected_word,
        inputs=[ai_pool_indices_state],
        outputs=[ai_export_status, ai_export_file],
    )

    preview_btn.click(
        fn=update_preview,
        inputs=[stem_box, options_box, answer_box, category_box, remark_box, image_path_box],
        outputs=[json_preview, extracted_state],
    )

    save_btn.click(
        fn=save_current_question,
        inputs=[stem_box, options_box, answer_box, category_box, remark_box, image_path_box],
        outputs=[save_status, json_preview, extracted_state],
    )

    clear_btn.click(
        fn=clear_entry,
        inputs=[],
        outputs=[
            question_input,
            figure_input,
            solution_input,
            stem_box,
            options_box,
            answer_box,
            category_box,
            remark_box,
            image_path_box,
            raw_output,
            parsed_json,
            json_preview,
        ],
    )

    search_btn.click(
        fn=search_database,
        inputs=[keyword_box, db_limit],
        outputs=[db_table, db_status, result_checks, search_result_choices_state],
    )
    select_all_btn.click(
        fn=select_all_search_results,
        inputs=[search_result_choices_state],
        outputs=[result_checks, db_status],
    )
    clear_result_checks_btn.click(
        fn=clear_search_checks,
        inputs=[],
        outputs=[result_checks, db_status],
    )
    add_selected_btn.click(
        fn=add_checked_to_selection,
        inputs=[result_checks, selected_indices_state],
        outputs=[selected_indices_state, selected_table, remove_checks, selected_status],
    )
    remove_selected_btn.click(
        fn=remove_checked_from_selection,
        inputs=[remove_checks, selected_indices_state],
        outputs=[selected_indices_state, selected_table, remove_checks, selected_status],
    )
    clear_selected_btn.click(
        fn=clear_selection,
        inputs=[],
        outputs=[selected_indices_state, selected_table, remove_checks, selected_status],
    )
    export_selected_btn.click(
        fn=export_selected_word,
        inputs=[selected_indices_state],
        outputs=[export_status, export_file],
    )
    export_all_btn.click(fn=export_all_word, inputs=[], outputs=[export_status, export_file])


if __name__ == "__main__":
    # For the packaged exe, bind only to localhost and open the browser automatically.
    port = int(os.environ.get("QUIZ_BUILDER_PORT", "7860"))
    url = f"http://localhost:{port}"

    def _open_browser_later():
        time.sleep(1.5)
        try:
            webbrowser.open(url)
        except Exception as e:
            print(f"[App] Failed to open browser automatically: {e}")
            print(f"[App] Please open manually: {url}")

    if os.environ.get("QUIZ_BUILDER_NO_BROWSER", "0") != "1":
        Thread(target=_open_browser_later, daemon=True).start()

    print(f"[App] Starting local server at {url}")
    demo.queue().launch(
        server_name="127.0.0.1",
        server_port=port,
        share=False,
        inbrowser=False,
        allowed_paths=[
        str(APP_DIR),
        str(EXPORT_DIR),
        str(IMAGE_DIR),],
        theme=gr.themes.Soft(primary_hue="violet", neutral_hue="slate", radius_size="lg"),
        css=CUSTOM_CSS,
    )
